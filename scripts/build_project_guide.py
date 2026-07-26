from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
ASSET_DIR = OUT_DIR / "assets"
OUT_PATH = OUT_DIR / "KSP-ATHENA_Complete_Project_Handbook.docx"

YELLOW = "FFD800"
BLACK = "111111"
CHARCOAL = "292929"
GRAY = "E9E9E9"
MID_GRAY = "666666"
WHITE = "FFFFFF"
CYAN = "22D3EE"
GREEN = "18B966"
RED = "C62828"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BLACK, size=8) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_text(cell, text: str, *, bold=False, color=BLACK, size=9.2) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    set_cell_margins(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_fixed_cell_width(cell, width_inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("PAGE ")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.bold = True
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_toc_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r' TOC \o "1-3" \h \z \u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and choose Update Field if page numbers are not visible."
    separate.append(placeholder)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def set_update_fields(document: Document) -> None:
    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(BLACK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.16

    for name, size, before, after in (
        ("Title", 30, 0, 14),
        ("Subtitle", 15, 0, 10),
        ("Heading 1", 18, 18, 8),
        ("Heading 2", 13.5, 13, 5),
        ("Heading 3", 11.5, 9, 3),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(BLACK)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    styles["Heading 1"].font.all_caps = True
    styles["Heading 2"].font.color.rgb = rgb(CHARCOAL)
    styles["Heading 3"].font.color.rgb = rgb(MID_GRAY)

    for name in ("Athena Eyebrow", "Athena Callout", "Athena Code", "Athena Small"):
        if name not in styles:
            styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

    eyebrow = styles["Athena Eyebrow"]
    eyebrow.font.name = "Calibri"
    eyebrow.font.size = Pt(8)
    eyebrow.font.bold = True
    eyebrow.font.all_caps = True
    eyebrow.font.color.rgb = rgb(MID_GRAY)
    eyebrow.paragraph_format.space_after = Pt(3)
    eyebrow.paragraph_format.keep_with_next = True

    callout = styles["Athena Callout"]
    callout.font.name = "Calibri"
    callout.font.size = Pt(10)
    callout.font.bold = True
    callout.font.color.rgb = rgb(BLACK)
    callout.paragraph_format.space_after = Pt(0)
    callout.paragraph_format.line_spacing = 1.1

    code = styles["Athena Code"]
    code.font.name = "Consolas"
    code.font.size = Pt(8.3)
    code.font.color.rgb = rgb(CHARCOAL)
    code.paragraph_format.space_after = Pt(4)
    code.paragraph_format.left_indent = Inches(0.18)

    small = styles["Athena Small"]
    small.font.name = "Calibri"
    small.font.size = Pt(8.5)
    small.font.color.rgb = rgb(MID_GRAY)
    small.paragraph_format.space_after = Pt(4)


def configure_section(section, *, title_page=False) -> None:
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.67)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.3)
    section.different_first_page_header_footer = title_page


class Guide:
    def __init__(self, doc: Document):
        self.doc = doc

    def p(self, text: str = "", *, bold=False, italic=False, style=None, align=None, keep=False):
        p = self.doc.add_paragraph(style=style)
        if align is not None:
            p.alignment = align
        p.paragraph_format.keep_together = keep
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        return p

    def h1(self, text: str, *, page_break=False):
        if page_break:
            self.doc.add_page_break()
        p = self.doc.add_heading(text, level=1)
        p.paragraph_format.keep_with_next = True
        # Yellow rule makes navigation visually match the application.
        p_pr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "22")
        bottom.set(qn("w:space"), "5")
        bottom.set(qn("w:color"), YELLOW)
        pbdr.append(bottom)
        p_pr.append(pbdr)
        return p

    def h2(self, text: str):
        return self.doc.add_heading(text, level=2)

    def h3(self, text: str):
        return self.doc.add_heading(text, level=3)

    def bullets(self, items: Iterable[str], level=0):
        for item in items:
            p = self.doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
            p.paragraph_format.space_after = Pt(3)
            p.add_run(item)

    def numbered(self, items: Iterable[str]):
        for item in items:
            p = self.doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(3)
            p.add_run(item)

    def callout(self, title: str, body: str, *, fill=YELLOW, label_color=BLACK):
        table = self.doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        set_fixed_cell_width(table.cell(0, 0), 1.55)
        set_fixed_cell_width(table.cell(0, 1), 5.85)
        set_cell_shading(table.cell(0, 0), fill)
        set_cell_shading(table.cell(0, 1), "F7F7F7")
        set_cell_text(table.cell(0, 0), title.upper(), bold=True, color=label_color, size=8.5)
        set_cell_text(table.cell(0, 1), body, size=9.5)
        set_table_borders(table, BLACK, 12)
        prevent_row_split(table.rows[0])
        self.doc.add_paragraph().paragraph_format.space_after = Pt(0)
        return table

    def table(self, headers: Sequence[str], rows: Sequence[Sequence[str]], widths=None, font_size=8.5):
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        hdr = table.rows[0]
        set_repeat_table_header(hdr)
        prevent_row_split(hdr)
        for i, header in enumerate(headers):
            set_cell_shading(hdr.cells[i], BLACK)
            set_cell_text(hdr.cells[i], header, bold=True, color=WHITE, size=8.3)
            if widths:
                set_fixed_cell_width(hdr.cells[i], widths[i])
        for row_values in rows:
            row = table.add_row()
            prevent_row_split(row)
            for i, value in enumerate(row_values):
                set_cell_shading(row.cells[i], WHITE if len(table.rows) % 2 else "F3F3F3")
                set_cell_text(row.cells[i], value, size=font_size)
                if widths:
                    set_fixed_cell_width(row.cells[i], widths[i])
        set_table_borders(table, BLACK, 8)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(0)
        return table

    def qa(self, question: str, answer: str):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(f"Q. {question}")
        run.bold = True
        run.font.color.rgb = rgb(BLACK)
        p2 = self.doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(5)
        label = p2.add_run("A. ")
        label.bold = True
        label.font.color.rgb = rgb(YELLOW)
        p2.add_run(answer)

    def code(self, text: str):
        p = self.doc.add_paragraph(style="Athena Code")
        p.add_run(text)
        p_pr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F2F2F2")
        p_pr.append(shd)
        return p

    def picture(self, path: Path, *, width: float, alt_text: str):
        shape = self.doc.add_picture(str(path), width=Inches(width))
        shape._inline.docPr.set("descr", alt_text)
        shape._inline.docPr.set("title", alt_text.split(".")[0][:120])
        return shape


def font(size: int, bold=False):
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def rounded_box(draw, xy, fill, outline=BLACK, width=4, radius=18):
    draw.rounded_rectangle(xy, radius=radius, fill=f"#{fill}", outline=f"#{outline}", width=width)


def centered_text(draw, box, text, fnt, fill=BLACK, spacing=5):
    x1, y1, x2, y2 = box
    bbox = draw.multiline_textbbox((0, 0), text, font=fnt, align="center", spacing=spacing)
    w, h = bbox[2] - bbox[0], bbox[3] - bbox[1]
    draw.multiline_text(((x1 + x2 - w) / 2, (y1 + y2 - h) / 2), text, font=fnt, fill=f"#{fill}", align="center", spacing=spacing)


def arrow(draw, start, end, color=BLACK, width=5):
    draw.line([start, end], fill=f"#{color}", width=width)
    x, y = end
    if abs(end[0] - start[0]) > abs(end[1] - start[1]):
        sign = 1 if end[0] > start[0] else -1
        points = [(x, y), (x - sign * 14, y - 9), (x - sign * 14, y + 9)]
    else:
        sign = 1 if end[1] > start[1] else -1
        points = [(x, y), (x - 9, y - sign * 14), (x + 9, y - sign * 14)]
    draw.polygon(points, fill=f"#{color}")


def make_architecture_diagram(path: Path) -> None:
    image = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(image)
    title = font(48, True)
    body = font(23, True)
    small = font(22, False)
    draw.text((70, 45), "KSP-ATHENA ACTIVE ARCHITECTURE", font=title, fill=f"#{BLACK}")
    draw.rectangle((70, 112, 1530, 124), fill=f"#{YELLOW}")

    boxes = {
        "browser": (80, 205, 430, 405),
        "catalyst": (625, 180, 980, 430),
        "data": (1160, 180, 1510, 430),
        "groq": (625, 585, 950, 790),
        "sarvam": (1160, 585, 1510, 790),
    }
    rounded_box(draw, boxes["browser"], "F7F7F7")
    rounded_box(draw, boxes["catalyst"], YELLOW)
    rounded_box(draw, boxes["data"], "F7F7F7")
    rounded_box(draw, boxes["groq"], CYAN)
    rounded_box(draw, boxes["sarvam"], CYAN)
    centered_text(draw, boxes["browser"], "WEB CLIENT\nReact + TypeScript + Vite", body)
    centered_text(draw, boxes["catalyst"], "ZOHO CATALYST\nAdvanced I/O Node function\nAuth · policy · analytics · APIs", body)
    centered_text(draw, boxes["data"], "CATALYST DATA STORE\nCaseRegistration\nAthenaUsers\nConversationHistory", body)
    centered_text(draw, boxes["groq"], "GROQ\nNL intent\nGrounded answers\nInterrogation assistance", body)
    centered_text(draw, boxes["sarvam"], "SARVAM AI\nSaaras STT · Sarvam-30B\nBulbul TTS", body)
    arrow(draw, (430, 305), (625, 305))
    arrow(draw, (980, 305), (1160, 305))
    arrow(draw, (790, 430), (790, 585))
    arrow(draw, (980, 370), (1250, 585))
    draw.text((455, 270), "HTTPS / JSON", font=small, fill=f"#{MID_GRAY}")
    draw.text((1000, 270), "ZCQL + SDK", font=small, fill=f"#{MID_GRAY}")
    draw.text((825, 485), "Server-side only", font=small, fill=f"#{MID_GRAY}")
    draw.text((1040, 495), "Voice + case Q&A", font=small, fill=f"#{MID_GRAY}")
    draw.text((80, 835), "Trust boundary: browser never receives provider API keys and never constructs executable database queries.", font=small, fill=f"#{BLACK}")
    image.save(path)


def make_query_pipeline(path: Path) -> None:
    image = Image.new("RGB", (1600, 920), "white")
    draw = ImageDraw.Draw(image)
    title = font(46, True)
    body = font(21, True)
    small = font(21, False)
    draw.text((65, 38), "CONVERSATIONAL QUERY: CONTROLLED DATA FLOW", font=title, fill=f"#{BLACK}")
    draw.rectangle((65, 105, 1535, 117), fill=f"#{YELLOW}")
    xs = [60, 330, 600, 870, 1140]
    labels = [
        ("1", "Natural-language\nquestion"),
        ("2", "Groq returns\nJSON intent only"),
        ("3", "Policy validates\nrole + filters"),
        ("4", "Server builds\nbounded SELECT"),
        ("5", "Grounded answer\n+ evidence panel"),
    ]
    for i, (num, label) in enumerate(labels):
        box = (xs[i], 245, xs[i] + 220, 485)
        rounded_box(draw, box, YELLOW if i in (2, 3) else "F7F7F7")
        draw.ellipse((xs[i] + 76, 165, xs[i] + 144, 233), fill=f"#{BLACK}")
        centered_text(draw, (xs[i] + 76, 165, xs[i] + 144, 233), num, font(30, True), WHITE)
        centered_text(draw, box, label, body)
        if i < 4:
            arrow(draw, (xs[i] + 220, 365), (xs[i] + 270, 365))
    guardrails = (160, 610, 1440, 835)
    rounded_box(draw, guardrails, "F2F2F2", width=3)
    draw.text((205, 640), "ENFORCED GUARDRAILS", font=body, fill=f"#{BLACK}")
    rules = [
        "Approved fields by role",
        "Approved filter keys and types",
        "Escaped text; wildcards rejected",
        "Maximum 5 conditions / 50 rows",
        "SELECT-only query constructed by server",
        "Answer model sees returned rows, not the database",
    ]
    for idx, rule in enumerate(rules):
        col = idx % 2
        row = idx // 2
        x = 215 + col * 610
        y = 695 + row * 42
        draw.rectangle((x, y + 5, x + 16, y + 21), fill=f"#{YELLOW}", outline=f"#{BLACK}", width=2)
        draw.text((x + 28, y), rule, font=small, fill=f"#{BLACK}")
    image.save(path)


def make_analytics_diagram(path: Path) -> None:
    image = Image.new("RGB", (1600, 920), "white")
    draw = ImageDraw.Draw(image)
    title = font(46, True)
    body = font(23, True)
    small = font(19, False)
    draw.text((65, 35), "RULE-BASED ANALYTICS AND HUMAN REVIEW", font=title, fill=f"#{BLACK}")
    draw.rectangle((65, 100, 1535, 112), fill=f"#{YELLOW}")
    center = (650, 310, 950, 510)
    rounded_box(draw, center, YELLOW)
    centered_text(draw, center, "NORMALIZED\nFIR RECORDS", font(30, True))
    modules = [
        ((80, 185, 400, 345), "Trends + social\ncomposition"),
        ((80, 565, 400, 725), "Hotspots +\nspatial surges"),
        ((1200, 185, 1520, 345), "Similar cases +\ncase brief"),
        ((1200, 565, 1520, 725), "Networks +\noffender profiles"),
        ((640, 690, 960, 850), "Forecast +\nearly warnings"),
    ]
    for box, text in modules:
        rounded_box(draw, box, "F4F4F4")
        centered_text(draw, box, text, body)
        if box[0] < 500:
            arrow(draw, (center[0], 410), (box[2], (box[1] + box[3]) // 2))
        elif box[0] > 1000:
            arrow(draw, (center[2], 410), (box[0], (box[1] + box[3]) // 2))
        else:
            arrow(draw, (800, center[3]), (800, box[1]))
    review = (450, 140, 1150, 230)
    rounded_box(draw, review, BLACK, outline=BLACK, width=2)
    centered_text(
        draw,
        review,
        "OUTPUT = PRIORITIZED REVIEW SIGNAL\nNOT A FINDING OF GUILT",
        font(20, True),
        WHITE,
    )
    draw.text((90, 870), "Every score retains source FIRs, contributing factors, coverage limits, and a non-causal limitation statement.", font=small, fill=f"#{BLACK}")
    image.save(path)


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    configure_section(section, title_page=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(44)
    run = p.add_run("KSP")
    run.font.name = "Calibri"
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = rgb(BLACK)
    run2 = p.add_run("–ATHENA")
    run2.font.name = "Calibri"
    run2.font.size = Pt(16)
    run2.font.bold = True
    run2.font.color.rgb = rgb(YELLOW)

    label = doc.add_paragraph(style="Athena Eyebrow")
    label.add_run("COMPLETE PROJECT REFERENCE • ARCHITECTURE • SECURITY • AI • FAQ")
    title = doc.add_paragraph(style="Title")
    title.add_run("KSP-ATHENA\nComplete Project Guide")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("A learning-first handbook for demonstrations, technical interviews, hackathon judging, and the path to production")

    bar = doc.add_table(rows=1, cols=1)
    bar.alignment = WD_TABLE_ALIGNMENT.LEFT
    bar.autofit = False
    set_fixed_cell_width(bar.cell(0, 0), 7.0)
    set_cell_shading(bar.cell(0, 0), YELLOW)
    set_cell_text(bar.cell(0, 0), "CONVERSATIONAL CRIME INTELLIGENCE • EXPLAINABLE ANALYTICS • HUMAN-IN-THE-LOOP DECISION SUPPORT", bold=True, size=11)
    set_table_borders(bar, BLACK, 14)

    doc.add_paragraph().paragraph_format.space_after = Pt(58)
    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.autofit = False
    rows = [
        ("Prepared for", "KSP-ATHENA project owner and hackathon team"),
        ("Edition", "1.0 — implementation snapshot"),
        ("Snapshot date", "26 July 2026"),
        ("Evidence basis", "Repository inspection and 109 passing automated tests"),
    ]
    for i, (key, value) in enumerate(rows):
        set_fixed_cell_width(meta.cell(i, 0), 1.55)
        set_fixed_cell_width(meta.cell(i, 1), 5.45)
        set_cell_shading(meta.cell(i, 0), BLACK)
        set_cell_shading(meta.cell(i, 1), "F3F3F3")
        set_cell_text(meta.cell(i, 0), key.upper(), bold=True, color=WHITE, size=8.5)
        set_cell_text(meta.cell(i, 1), value, size=9.5)
    set_table_borders(meta, BLACK, 8)
    doc.add_paragraph()
    note = doc.add_paragraph(style="Athena Small")
    note.add_run(
        "Scope note: this handbook distinguishes the actively deployed React + Zoho Catalyst implementation "
        "from legacy prototype folders that are not part of the current deployment."
    )


def add_headers_and_footers(doc: Document) -> None:
    for index, section in enumerate(doc.sections):
        if index == 0:
            continue
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("KSP-ATHENA  |  COMPLETE PROJECT GUIDE")
        r.font.name = "Calibri"
        r.font.size = Pt(8)
        r.font.bold = True
        r.font.color.rgb = rgb(MID_GRAY)
        p_pr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "10")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), YELLOW)
        pbdr.append(bottom)
        p_pr.append(pbdr)
        footer = section.footer
        fp = footer.paragraphs[0]
        add_page_number(fp)


def add_front_matter(g: Guide) -> None:
    g.h1("How to use this guide")
    g.p(
        "This is both a project reference and a compact learning course. Read Parts I–III before a demo, "
        "Parts IV–VIII before a technical discussion, and the FAQ before judging or interviews. The wording "
        "is deliberately honest: a system can be impressive as a hackathon prototype while still needing "
        "substantial engineering, governance, and validation before real police deployment."
    )
    g.callout(
        "Fast answer",
        "KSP-ATHENA is a working hackathon-grade, serverless crime-intelligence prototype. It combines controlled natural-language querying, case-scoped bilingual voice conversation, deterministic analytical modules, evidence-linked explanations, role-aware access, and audit logging. It is not yet a production law-enforcement system.",
    )
    g.h2("Reading paths")
    g.table(
        ["If you are preparing for…", "Read first", "Outcome"],
        [
            ("A 5-minute demo", "Executive summary, product tour, demo script", "You can narrate the value and show the strongest workflow."),
            ("A judge Q&A", "Architecture, AI safety, analytics, FAQ", "You can explain design choices and defend limitations."),
            ("A technical interview", "Stack, API, data model, tests, deployment", "You can discuss the project as an engineer, not only as a user."),
            ("Production planning", "Security, governance, gaps, roadmap, evaluations", "You can convert the prototype into a phased delivery plan."),
            ("Career learning", "Learning lessons, glossary, endpoint and role appendices", "You understand the concepts behind the implementation."),
        ],
        [1.65, 2.25, 3.1],
    )
    g.h2("Legend used throughout")
    g.table(
        ["Label", "Meaning"],
        [
            ("ACTIVE", "Implemented in the current frontend and Catalyst function."),
            ("PARTIAL", "Useful implementation exists, but the challenge requirement is not fully satisfied."),
            ("PROTOTYPE / DORMANT", "Code exists in the repository but is not wired into the active deployment."),
            ("ROADMAP", "Recommended future capability; do not claim it as currently implemented."),
            ("HUMAN REVIEW", "Output supports an authorized decision-maker; it is not an autonomous finding."),
        ],
        [1.55, 5.45],
    )
    g.h2("Table of contents")
    g.table(
        ["Part", "Coverage"],
        [
            ("I — Executive understanding", "Vision, problem, design principle, differentiators, and maturity"),
            ("II — Product and user experience", "Personas, feature catalogue, journey, and UI design"),
            ("III — Architecture and data flow", "Active deployment, trust boundaries, query and analytics pipelines"),
            ("IV — Technology stack", "Every active technology, reasons, trade-offs, and dormant prototypes"),
            ("V — Data model, API, workflows", "Tables, endpoints, authentication, FIR intake, synthetic data"),
            ("VI — AI, analytics, explainability", "Models, grounding, rules, scores, forecasts, evidence, guardrails, evals"),
            ("VII — Security, testing, deployment", "RBAC, audit, limitations, 109 tests, environment, release flow"),
            ("VIII — Challenge and roadmap", "Coverage of all ten challenge areas and phased production plan"),
            ("IX — Demo and learning", "Demo script, test prompts, recovery, and career lessons"),
            ("X — Detailed FAQ", "Product, architecture, AI, analytics, security, operations, and ethics Q&A"),
            ("XI — Appendices", "Repository map, configuration, glossary, checklists"),
        ],
        [2.45, 4.55],
        font_size=8.2,
    )


def add_executive_summary(g: Guide) -> None:
    g.h1("Part I — Executive understanding", page_break=True)
    g.h2("1. Project in one sentence")
    g.p(
        "KSP-ATHENA is an explainable, multilingual crime-intelligence workspace that lets authorized users "
        "query FIR records in natural language, examine one case deeply, discover recorded relationships and "
        "patterns, and receive evidence-linked decision-support signals through a secure serverless application."
    )
    g.h2("2. The problem it addresses")
    g.p(
        "Crime databases are usually structured around tables and rigid forms, while investigators think in "
        "questions: “Show recent vehicle thefts in Koramangala,” “Which FIRs mention the same accused?”, or "
        "“What evidence in this statement supports the claim that faces were not visible?” KSP-ATHENA reduces "
        "that translation burden. It also brings commonly separated tasks—retrieval, summarization, mapping, "
        "network review, trend analysis, forecasting, and evidence explanation—into one workspace."
    )
    g.h2("3. The central design idea")
    g.p(
        "The platform uses generative AI only where language flexibility is valuable, and uses deterministic "
        "server-side rules where reproducibility and control matter. A language model may translate a question "
        "into a small JSON search intent, but it cannot send SQL to the database. The server validates the intent, "
        "constructs a bounded SELECT query, executes it, and gives only the returned fields to the answer model. "
        "Analytical scores are calculated by visible code and include evidence, thresholds, coverage, and limitations."
    )
    g.callout(
        "Core principle",
        "Natural language is flexible; authority is not. The user can ask freely, but the server decides which fields, filters, rows, and analytical modules the authenticated role may access.",
    )
    g.h2("4. What makes the prototype distinctive")
    g.bullets(
        [
            "Case-scoped English/Kannada voice conversation: speech-to-text, selected-FIR-only reasoning, and spoken response.",
            "A controlled natural-language query pipeline that never executes model-generated SQL.",
            "A uniform evidence panel that cites FIR numbers, fields, filters, data source, record limits, and reasoning steps.",
            "Deterministic similar-case matching, offender review priority, network links, hotspot alerts, forecasts, and early warnings.",
            "A statewide synthetic dataset designed with repeat suspects, co-accused relations, shared indicators, varied regions, and long statements.",
            "A visually coherent yellow/black neo-brutalist interface with a reorganized vertical feature menu and loading states.",
            "Role-aware field access, signed sessions, salted password hashing, and signed audit events.",
        ]
    )
    g.h2("5. Current maturity")
    g.table(
        ["Dimension", "Current position", "Interpretation"],
        [
            ("Hackathon demonstration", "READY WITH CAVEATS", "The integrated workflows are suitable for a controlled demo using synthetic data."),
            ("Functional prototype", "STRONG", "The major modules are active and backed by automated tests."),
            ("Pilot with real police data", "NOT YET", "Requires security review, data agreements, identity integration, evaluations, observability, and operational controls."),
            ("Statewide production", "NOT YET", "Requires formal architecture, scale testing, governance, legal review, high availability, and ongoing model/data assurance."),
        ],
        [1.6, 1.4, 4.0],
    )


def add_personas_and_tour(g: Guide) -> None:
    g.h1("Part II — Product and user experience", page_break=True)
    g.h2("6. Intended users")
    g.table(
        ["Persona", "Primary needs", "How KSP-ATHENA helps"],
        [
            ("Constable", "Basic record lookup and FIR intake", "Restricted conversational fields, dashboard access, and validated FIR registration."),
            ("Investigator", "Case facts, statements, similar cases, links, leads", "Deep Dive, case conversation, briefs, network analysis, profiles, and early warnings."),
            ("Analyst", "Cross-case patterns, geography, trends, and forecasting", "Full analytical views and permitted sensitive investigation fields."),
            ("Supervisor", "Oversight, operational review, governance", "Full analytical access, audit trail, and administrative dataset controls."),
            ("Policymaker", "Aggregate patterns and preventive planning", "The design supports this persona, but a separate production-grade aggregate-only role is still needed."),
            ("Argos demo visitor", "Explore every feature without registration", "A 30-minute full-access demo session; useful for public demonstrations but unsafe as a production design."),
        ],
        [1.25, 2.25, 3.5],
    )
    g.h2("7. Why the demo role is called Argos")
    g.p(
        "Argos is a fitting name because it evokes watchfulness. In Greek mythology, Argus Panoptes was the "
        "all-seeing guardian; “Argos” therefore signals an observer that can look across the platform. In this "
        "implementation, however, Argos is more than read-only: following the project requirement, it receives "
        "full demo access and can register cases. The name is memorable, but the permissions must be redesigned "
        "before any public production deployment."
    )
    g.h2("8. Navigation and feature catalogue")
    feature_rows = [
        ("Dashboard", "Core", "ACTIVE", "Jurisdiction filter, status totals, crime classification breakdown, and basic case list for selected divisions."),
        ("Conversational Hub", "Core", "ACTIVE", "Natural-language FIR search, follow-up context, English/Kannada display, voice input, PDF history export, evidence panel."),
        ("Case Deep Dive", "Core", "ACTIVE", "Loads one FIR, generates brief/timeline/leads/similar cases, supports case report PDF and selected-case voice conversation."),
        ("Register Fresh FIR", "Core", "ACTIVE", "Validated write flow for victim, location, accused, crime type, incident date, and statement."),
        ("Early Warnings", "Intelligence", "ACTIVE", "Prioritizes spatial surge, repeat association, network activity, and aggregate forecast signals."),
        ("Criminal Network", "Intelligence", "ACTIVE", "Creates evidence-qualified accused-to-accused and accused-to-case graph links."),
        ("Offender Profiles", "Intelligence", "ACTIVE", "Groups recorded accused names and calculates a transparent review-priority score."),
        ("Trend Analytics", "Analytics", "ACTIVE", "6/12/24-month trends by month, crime type, division, and victim age band."),
        ("Spatial Alerts", "Analytics", "ACTIVE", "Leaflet hotspot map, coordinate coverage, mapping provenance, and 30-day concentration alerts."),
        ("Crime Forecast", "Analytics", "ACTIVE", "Three-month aggregate baseline with data sufficiency, backtest error, range, and filters."),
        ("Social Insights", "Analytics", "PARTIAL", "Descriptive victim-age, crime-type, and geography composition; no external socio-economic variables yet."),
        ("Audit Trail", "Governance", "ACTIVE", "Supervisor/Argos view of signed server-side operational events with filters."),
        ("Financial link analysis", "Challenge area", "PARTIAL", "Transaction/account text can appear in statements and network indicators, but no bank feed or money-trail model is integrated."),
    ]
    g.table(["Feature", "Group", "Status", "What it does"], feature_rows, [1.25, 1.0, 0.75, 4.0], font_size=8.1)
    g.h2("9. Recommended end-to-end user journey")
    g.numbered(
        [
            "Authenticate through registration/login or start the short-lived Argos demo.",
            "Open Dashboard to understand the current jurisdiction and data coverage.",
            "Ask a natural-language question in Conversational Hub and inspect the uniform evidence panel.",
            "Open Case Deep Dive for a cited FIR and review structured facts, statement, timeline, evidence leads, and similar cases.",
            "Use Talk About This Case in English or Kannada; keep the conversation scoped to the selected FIR.",
            "Review network, offender profile, hotspot, trend, social, forecast, and early-warning modules as appropriate.",
            "Download a one-page KSP-ATHENA case report or conversation PDF.",
            "For governance, use Audit Trail to verify that sensitive actions were recorded.",
        ]
    )
    g.h2("10. User-interface design language")
    g.p(
        "The interface uses a high-contrast neo-brutalist visual language: black borders and shadows, white cards, "
        "KSP yellow as the primary accent, and cyan for the Argos demo badge/button. Important modules—Dashboard, "
        "Conversational Hub, and Case Deep Dive—receive stronger visual treatment. Features are grouped vertically "
        "in a persistent left navigation panel, while Register Fresh FIR remains a prominent top-right action. "
        "The former continuously moving status bar was removed; loading animations now appear only during real work."
    )
    g.callout(
        "Accessibility note",
        "High contrast is useful, but visual accessibility is broader than contrast. Production work should add keyboard-flow testing, focus indicators, screen-reader labels, reduced-motion support, Kannada font verification, responsive mobile layouts, and WCAG 2.2 AA assessment.",
        fill=CYAN,
    )


def add_architecture(g: Guide, architecture_path: Path, query_path: Path, analytics_path: Path) -> None:
    g.h1("Part III — Architecture and data flow", page_break=True)
    g.h2("11. Active deployment architecture")
    g.p(
        "The active system is a serverless web application. The React frontend is compiled into static assets and "
        "served by Zoho Catalyst. It calls one Catalyst Advanced I/O Node.js function. That function owns authentication, "
        "authorization, validation, database access, analytical rules, audit writing, and calls to Groq and Sarvam AI."
    )
    g.picture(
        architecture_path,
        width=7.0,
        alt_text="Active KSP-ATHENA architecture. A React browser client calls a Zoho Catalyst Advanced I/O function, which accesses Catalyst Data Store and server-side Groq and Sarvam AI services.",
    )
    cap = g.doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run("Figure 1. Active KSP-ATHENA architecture and trust boundaries.")
    g.h2("12. Why this architecture was chosen")
    g.table(
        ["Choice", "Why it fits the prototype", "Trade-off"],
        [
            ("Single React client", "Fast iteration, rich interaction, reusable components, and deployable static bundle.", "Client-side state and localStorage are convenient but not sufficient for high-assurance sessions."),
            ("Single Advanced I/O function", "Centralizes security policy and keeps deployment simple.", "A large function can become a monolith; production should split domains or establish clear internal boundaries."),
            ("Catalyst Data Store", "Already provisioned, close to the function, and avoids operating a separate database.", "Schema evolution, query patterns, indexing, and portability need deliberate management."),
            ("External model APIs", "Provides strong language and speech capabilities without model hosting.", "Adds vendor dependency, privacy review, cost, availability, and network latency concerns."),
            ("Deterministic analytics in Node", "Easy to explain, test, and deploy with the function.", "Not a replacement for validated criminology models, geospatial platforms, or large-scale analytical infrastructure."),
        ],
        [1.55, 3.2, 2.25],
    )
    g.h2("13. Conversational query pipeline")
    g.picture(
        query_path,
        width=7.0,
        alt_text="Controlled conversational query pipeline. Natural language becomes JSON intent, server policy validates role and filters, the server builds a bounded SELECT, and the answer includes an evidence panel.",
    )
    cap = g.doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run("Figure 2. The model proposes an intent; only server policy can authorize and execute a lookup.")
    g.numbered(
        [
            "The authenticated user asks a natural-language question.",
            "Recent messages from the same user and session are loaded to resolve follow-up context.",
            "Groq is prompted to return a JSON search intent, not SQL or executable text.",
            "The server rejects unknown keys, unauthorized fields, unsupported filters, invalid values, wildcards, excessive limits, or more than five conditions.",
            "The server constructs a SELECT-only ZCQL query with an explicit field list and a limit of 1–50 rows.",
            "Catalyst Data Store returns permitted fields only.",
            "Groq synthesizes an answer from those rows; it does not receive database credentials or arbitrary database access.",
            "The server generates the evidence contract independently: cited FIRs, returned fields, filters, data source, result bound, reasoning steps, and limitations.",
            "The user and assistant turns are stored in ConversationHistory for contextual follow-up and local PDF export.",
        ]
    )
    g.h2("14. Analytics pipeline")
    g.picture(
        analytics_path,
        width=7.0,
        alt_text="Rule-based analytics flow. Normalized FIR records feed trends, social composition, hotspots, similar cases, case briefs, networks, profiles, forecasts, and early-warning review signals.",
    )
    cap = g.doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run("Figure 3. Deterministic analytical outputs are framed as review signals.")
    g.p(
        "Most analytical modules retrieve a bounded set of up to 300 recent CaseRegistration rows, normalize fields, "
        "apply deterministic rules, and return both results and coverage metadata. This makes a demonstration responsive "
        "and explainable, but the 300-row cap is a material scalability limit. Production analytics should use pagination, "
        "pre-aggregations, an analytical store, or carefully indexed database queries rather than loading a fixed sample."
    )
    g.h2("15. Trust boundaries")
    g.table(
        ["Boundary", "What is trusted", "What is treated as untrusted"],
        [
            ("Browser → function", "Signed session after server verification", "Every body, query string, file, localStorage value, and role claim from the client."),
            ("Function → database", "Server-owned field policy and constructed query", "Natural-language text and model output until validated."),
            ("Function → Groq", "Bounded prompt and returned database rows", "Model output; it must pass JSON/policy validation where used for intent."),
            ("Function → Sarvam", "Selected role-permitted FIR context", "FIR statement and conversation content are evidence, never instructions."),
            ("Analytics → investigator", "Recorded fields and deterministic calculations", "Identity matches, correlations, priority scores, and forecasts as final truth."),
        ],
        [1.35, 2.45, 3.2],
    )


def add_stack(g: Guide) -> None:
    g.h1("Part IV — Technology stack and the reason for every choice", page_break=True)
    g.h2("16. Active frontend stack")
    frontend_rows = [
        ("React 19.2.8", "Component-based web UI", "Supports modular feature screens, stateful interactions, and reusable evidence/visual components.", "A mature ecosystem; requires discipline around state, effects, and accessibility."),
        ("TypeScript 5.9.3", "Static typing", "Catches interface and payload mistakes before deployment and improves maintainability.", "Types do not validate runtime network data by themselves."),
        ("Vite 6.4.3", "Development server and bundler", "Fast local feedback and compact production build with straightforward configuration.", "Production behavior still depends on hosting paths and browser compatibility."),
        ("Leaflet 1.9.4 + React Leaflet 5.0.0", "Interactive maps", "Widely used, open mapping primitives and natural React integration.", "Needs licensed tile strategy, clustering, geocoding, and large-data optimization in production."),
        ("Recharts 3.10.0", "Charts", "Declarative React charts for trends, composition, diagnostics, and forecasts.", "Must ensure accessible text equivalents and correct interpretation."),
        ("react-force-graph-2d 1.29.1", "Network visualization", "Interactive graph exploration for people-to-case relationships.", "Dense graphs require layout controls, filtering, and performance tuning."),
        ("Lucide React 1.26.0", "Icons", "Consistent lightweight icon set that improves feature recognition.", "Icons need labels; they cannot carry meaning alone."),
        ("html2canvas 1.4.1", "DOM capture", "Allows browser-rendered case and conversation layouts to become report images.", "Raster capture can reduce selectable text and accessibility."),
        ("jsPDF 4.2.1", "Client PDF creation", "Enables local one-page case reports and conversation exports without server storage.", "Long content, font embedding, pagination, and evidentiary authenticity need stronger production handling."),
    ]
    g.table(["Technology", "Role", "Why selected", "Important trade-off"], frontend_rows, [1.42, 1.25, 2.65, 1.68], font_size=7.9)
    g.h2("17. Active backend and cloud stack")
    backend_rows = [
        ("Zoho Catalyst", "Serverless platform", "Hosts the static client, function, environment variables, and Data Store in one project.", "Platform-specific deployment and ZCQL; needs production configuration governance."),
        ("Advanced I/O Function", "HTTP API runtime", "Allows a conventional Express application and multiple endpoints in one Catalyst function.", "Cold starts, execution limits, and monolith growth must be monitored."),
        ("Node.js / CommonJS", "Backend runtime", "Matches the JavaScript frontend skill set and suits API orchestration and deterministic rules.", "CPU-heavy analytics may need separate services."),
        ("Express 5.2.1", "Routing and middleware", "Clear endpoint structure, JSON handling, authentication middleware, and errors.", "Security headers, rate limiting, and structured validation need explicit additions."),
        ("zcatalyst-sdk-node", "Catalyst integration", "Provides Data Store/ZCQL access inside the function.", "Pinned versions are safer than `latest` for reproducible production builds."),
        ("Groq SDK 1.3.0", "LLM access", "Low-latency inference for intent conversion and grounded answer/interrogation prompts.", "Third-party processing, model lifecycle, quotas, and hallucination risk."),
        ("Axios 1.18.1", "HTTP client", "Simple calls to Sarvam APIs with headers and binary responses.", "Timeout, retries, circuit breaking, and telemetry need systematic configuration."),
        ("Multer 2.2.0", "Audio upload", "Parses multipart voice recordings for speech-to-text.", "Must enforce strict type/size scanning and avoid unnecessary retention."),
        ("Form-Data 4.0.6", "Multipart forwarding", "Builds the Sarvam speech-to-text request.", "Boundary handling and stream errors must be controlled."),
        ("dotenv 17.4.2", "Local environment loading", "Keeps local secrets outside source code.", "Production must use Catalyst encrypted variables, not committed `.env` files."),
        ("Node crypto", "Security primitives", "Implements scrypt password hashes, HMAC-signed sessions, stable pseudonymous IDs, random IDs, and audit signatures.", "Custom token formats demand careful review; standard identity platforms are preferred in production."),
    ]
    g.table(["Technology", "Role", "Why selected", "Important trade-off"], backend_rows, [1.42, 1.25, 2.65, 1.68], font_size=7.85)
    g.h2("18. AI services")
    g.table(
        ["Service / model", "Active use", "Reason", "Boundary"],
        [
            ("Groq / llama-3.1-8b-instant", "Natural-language → JSON search intent; grounded answer synthesis; statement interrogation", "Responsive language handling suitable for a live demo.", "No raw SQL execution; answer must be grounded in returned data or selected statement evidence."),
            ("Sarvam AI / Saaras v3", "Speech-to-text for English/Kannada questions", "India-focused speech capability and regional-language support.", "Audio is sent to a third party; consent, retention, encryption, and data residency require review."),
            ("Sarvam AI / Sarvam-30B", "Selected-FIR-only bilingual case conversation", "Natural bilingual responses from a tightly bounded case context.", "This is prompt grounding, not model fine-tuning; the output remains probabilistic."),
            ("Sarvam AI / Bulbul v3", "Text-to-speech response", "Completes the two-way spoken interaction.", "Voice output should never be treated as an official legal statement without verification."),
        ],
        [1.55, 2.0, 1.9, 1.55],
        font_size=8.0,
    )
    g.callout(
        "Terminology",
        "KSP-ATHENA does not currently fine-tune Sarvam-30B. It performs case-scoped retrieval/prompt grounding: selected FIR fields are placed in the model context. Fine-tuning would mean training model weights on a curated dataset, with a separate evaluation and governance program.",
        fill=CYAN,
    )
    g.h2("19. Testing and development stack")
    g.table(
        ["Technology", "Purpose", "Why it matters"],
        [
            ("Node built-in test runner", "86 backend unit and policy tests", "No extra runner is required; TAP output is simple and fast."),
            ("Vitest 4.1.10", "23 frontend tests across 6 files", "Integrates naturally with Vite and supports fast component testing."),
            ("React Testing Library 16.3.2", "User-oriented component tests", "Encourages testing behavior and accessible interactions instead of implementation internals."),
            ("jest-dom 6.9.1", "Readable DOM assertions", "Improves clarity of UI expectations."),
            ("jsdom 26.1.0", "Browser-like test environment", "Allows React components to run in Node; it is not a real browser and cannot replace end-to-end testing."),
            ("TypeScript `tsc --noEmit`", "Build-time type check", "Blocks production bundling when types do not compile."),
        ],
        [1.85, 2.05, 3.1],
    )
    g.h2("20. Dormant prototype stack—do not present as deployed")
    g.p(
        "The repository also contains earlier or exploratory backend folders. They are useful evidence of design exploration, "
        "but `catalyst.json` deploys only `frontend/build` and `functions/ks_intelli_pol_function`. Therefore the following "
        "must not be described as active runtime components:"
    )
    g.table(
        ["Folder", "Technology / idea", "Current status"],
        [
            ("backend/ml-services-python", "Python, FastAPI-style services, pandas, scikit-learn, NetworkX, python-louvain", "Prototype / disconnected; not called by active routes."),
            ("backend/python-analytics", "Python analytical service", "Prototype / disconnected."),
            ("backend/orchestration-node", "Alternative Node orchestration, PDFKit, RBAC, voice bridge, audit modules", "Legacy / disconnected."),
            ("backend/agents", "Experimental SQL and citation agents", "Legacy / disconnected; the active safe query policy lives in the Catalyst function."),
            ("config/prompts", "Early NL-to-SQL prompt assets", "Reference only; the active backend uses a JSON-intent pipeline."),
        ],
        [2.05, 2.75, 2.2],
    )
    g.p(
        "Career lesson: repository code is not the same as deployed architecture. Always trace the build and deployment "
        "configuration, imports, network calls, and runtime entry points before claiming a technology is in production."
    )


def add_data_model(g: Guide) -> None:
    g.h1("Part V — Data model, API, and workflows", page_break=True)
    g.h2("21. Core tables")
    g.h3("21.1 CaseRegistration")
    g.p(
        "This manually provisioned Catalyst Data Store table is the core operational dataset. It stores the FIR identifier, "
        "victim and accused fields, crime classification, statement, location/jurisdiction, status, time, and optional coordinates. "
        "The application assumes the following fields across its routes."
    )
    g.table(
        ["Field", "Use in the application", "Sensitivity / caveat"],
        [
            ("CrimeNo", "Primary case reference and citation", "Must be unique; currently a random six-digit number with collision retry."),
            ("VictimName / VictimAge / VictimMobile / VictimAddress", "Case detail, investigation search, social composition", "Personal data; role-restricted and should be encrypted/masked by policy."),
            ("AccusedName / AccusedAge / AccusedMobile", "Case detail, profiles, network links", "An accusation is not guilt; identity matching requires independent verification."),
            ("Pincode / DivisionName", "Dashboard, filtering, trends, maps, networks", "Data normalization is required; current records may use City/Division suffix variants."),
            ("CrimeTypeID / CrimeTypeName", "Classification, retrieval, trends, forecast", "Production needs controlled taxonomy and versioning."),
            ("VictimStatement", "Deep Dive, case conversation, evidence extraction, method indicators", "Highly sensitive free text and a prompt-injection surface."),
            ("CaseStatus", "Dashboard, retrieval, active-case calculations", "Needs controlled workflow and authoritative status source."),
            ("RegisteredAt / RegisteredBy", "Time analysis, provenance, audit context", "Synthetic records are labeled by seed version."),
            ("Latitude / longitude", "Direct hotspot coordinates", "Coverage is incomplete; fallback mapping is disclosed."),
        ],
        [1.7, 3.1, 2.2],
        font_size=8.0,
    )
    g.h3("21.2 AthenaUsers")
    g.table(
        ["Column", "Purpose", "Control"],
        [
            ("Username", "Unique sign-in identifier", "3–50 characters; letters, numbers, dot, underscore, or hyphen."),
            ("PasswordHash", "scrypt-derived password representation", "Password is never stored in plaintext."),
            ("PasswordSalt", "Random per-user salt", "Prevents identical passwords from producing identical stored values."),
            ("Role", "Constable, Investigator, Supervisor, or Analyst", "Currently user-selected during registration with a shared four-digit code—a demo compromise."),
            ("IsActive", "Account enable/disable flag", "Checked during login."),
            ("CreatedAt", "Account creation time", "Supports administration and audit context."),
        ],
        [1.45, 2.65, 2.9],
    )
    g.h3("21.3 ConversationHistory")
    g.table(
        ["Column", "Conversation use", "Audit use"],
        [
            ("ConversationID", "Random application message ID", "Random event ID in a high numeric range."),
            ("SessionID", "Groups conversation turns", "Mirrors event ID."),
            ("UserID", "HMAC-derived pseudonymous user ID", "Stable pseudonymous actor ID."),
            ("Role", "`user` or `assistant`", "Authenticated role."),
            ("Content", "JSON: text, mode, citations, explainability", "JSON: action, outcome, target, safe details, timestamp, HMAC signature."),
            ("Language", "`en` or `kn`", "`audit-v1` marker isolates audit records."),
            ("CREATEDTIME", "System insertion timestamp", "Ordering for the audit view."),
        ],
        [1.35, 2.82, 2.83],
    )
    g.callout(
        "Schema decision",
        "Audit events reuse ConversationHistory because the originally expected ConversationLog table was not provisioned. The `audit-v1` language marker separates event rows from ordinary chat rows. This unblocked the hackathon, but a dedicated append-only audit store is preferable for production.",
    )
    g.h2("22. Active HTTP API reference")
    routes = [
        ("POST", "/auth/register", "Public", "Create a user after username/password/role/shared-code validation."),
        ("POST", "/auth/login", "Public", "Verify scrypt password and return an 8-hour signed session."),
        ("POST", "/auth/demo", "Public", "Return a 30-minute signed Argos session."),
        ("POST", "/chat", "Authenticated", "Controlled conversational search and grounded answer."),
        ("GET", "/conversation-history", "Authenticated", "Load one user/session’s bounded message history."),
        ("GET", "/audit-events", "Supervisor, Argos", "Read and verify signed audit events; optional filters."),
        ("POST", "/register", "Authenticated", "Validate and insert one fresh FIR."),
        ("POST", "/admin/seed-synthetic-firs", "Supervisor", "Idempotently load labeled synthetic datasets with confirmation."),
        ("POST", "/fetch-case", "Authenticated", "Load one FIR with role-permitted fields."),
        ("POST", "/case-intelligence", "Authenticated", "Return similar cases and permitted repeat associations."),
        ("POST", "/case-brief", "Authenticated", "Return extractive overview, facts, timeline, and evidence leads."),
        ("POST", "/case-conversation", "Authenticated", "Selected-FIR-only Sarvam response and optional audio."),
        ("POST", "/interrogate", "Authenticated", "Statement-grounded question answering with evidence candidates."),
        ("GET", "/dashboard-metrics", "Authenticated", "Jurisdiction totals, breakdown, divisions, and selected case list."),
        ("GET", "/trend-analytics", "Authenticated", "6/12/24-month descriptive trends."),
        ("GET", "/spatial-hotspots", "Authenticated", "Hotspots, mapping coverage, and spatial alerts."),
        ("GET", "/offender-profiles", "Investigation roles", "Explainable name-grouped review profiles."),
        ("GET", "/sociological-insights", "Authenticated", "Descriptive age/crime/division composition."),
        ("GET", "/crime-forecast", "Authenticated", "Aggregate three-month transparent forecast."),
        ("GET", "/criminal-network", "Investigation roles", "Evidence-qualified associations and graph."),
        ("GET", "/early-warnings", "Investigation roles", "Prioritized compound review alerts."),
        ("POST", "/transcribe", "Authenticated", "Forward bounded audio to Sarvam Saaras v3."),
    ]
    g.table(["Method", "Route", "Access", "Responsibility"], routes, [0.72, 1.75, 1.28, 3.25], font_size=7.7)
    g.h2("23. Authentication lifecycle")
    g.numbered(
        [
            "Registration validates the username, password, selected role, and registration code. The default code is 3024 and can be changed with `ATHENA_REGISTRATION_CODE`.",
            "The server generates a random salt and uses scrypt to derive the password hash.",
            "Login finds the user, checks IsActive, re-derives the hash, and compares it safely.",
            "The server returns an HMAC-signed token containing the username, role, issue time, and expiry.",
            "The browser stores token, role, and username in localStorage and sends the token in `X-Athena-Token` (Bearer compatibility also exists).",
            "Protected middleware verifies signature and expiry before assigning server-trusted identity to the request.",
            "Normal sessions last eight hours; Argos demo sessions last thirty minutes.",
            "A 401 response clears local browser session values and redirects to login.",
        ]
    )
    g.h2("24. FIR registration workflow")
    g.p(
        "The active form accepts victim name/age/mobile, location, pincode, optional accused name, one of five crime types, "
        "incident date, and a 10–5000 character statement. The server rejects unknown fields, invalid Indian phone numbers, "
        "invalid pincodes, out-of-range age, unsupported crime type, control characters, and future dates. It inserts the "
        "record with status Registered and retries random six-digit CrimeNo allocation on collisions."
    )
    g.callout(
        "Known gap",
        "The registration payload currently does not derive or insert DivisionName from the Karnataka pincode map, and it stores the incident date at midnight as RegisteredAt. Production should separate IncidentAt from RegisteredAt, resolve jurisdiction server-side, and preserve the authentic registration timestamp.",
        fill=CYAN,
    )
    g.h2("25. Synthetic dataset")
    g.p(
        "Two labeled, idempotent seed sets exist: V1 contains 25 FIRs (926001–926025) and V2 contains 30 FIRs "
        "(927001–927030). They cover the supported crime types, multiple Karnataka divisions and time periods, long "
        "statements, repeat suspect names, co-accused relations, financial references, hotspots, and selected links to "
        "earlier victims. The loader inserts only missing numbers and reports conflicts instead of overwriting rows."
    )
    g.bullets(
        [
            "All synthetic statements are explicitly labeled as test records.",
            "Synthetic data is for feature demonstration and unit tests; it cannot validate real-world accuracy or fairness.",
            "The UI seed box was removed after loading, but the supervisor-only administrative endpoint and files remain.",
            "Never mix synthetic records into an operational crime database without strong environment and provenance separation.",
        ]
    )


def add_ai_and_analytics(g: Guide) -> None:
    g.h1("Part VI — AI, analytics, explainability, and guardrails", page_break=True)
    g.h2("26. Conversational Hub in detail")
    g.h3("26.1 Intent generation")
    g.p(
        "The first Groq call uses `llama-3.1-8b-instant` at temperature 0 to convert the question and bounded context into "
        "JSON containing `requestedFields`, `filters`, `sort`, and `limit`. Raw SQL, code, unknown properties, and additional "
        "operators are forbidden by prompt and rejected by code."
    )
    g.h3("26.2 Policy validation")
    g.p(
        "The query policy maintains server-side allowlists for fields and filters by role. It validates object shape, rejects "
        "unknown keys, checks integer and date bounds, normalizes permitted contains filters, escapes quote characters, rejects "
        "user wildcards, permits at most five conditions, limits rows to 1–50, and constructs the final ZCQL SELECT."
    )
    g.h3("26.3 Grounded answer synthesis")
    g.p(
        "The second Groq call uses temperature 0.3 and receives only the permitted returned records plus conversational context. "
        "The evidence panel is not trusted to the model: the server independently derives citations and reasoning metadata from "
        "the validated intent, executed query metadata, and returned rows."
    )
    g.h2("27. Uniform Explainable-AI Evidence Panel")
    g.table(
        ["Element", "Meaning", "Why it matters"],
        [
            ("Evidence status", "No matches, matching records, or bounded result limit reached", "Prevents an empty result from being confused with a database-wide proof of absence."),
            ("FIR citations", "Crime numbers of returned records", "Lets the investigator open primary records."),
            ("Fields used", "The exact permitted fields returned", "Shows what the answer could and could not know."),
            ("Filters applied", "Human-readable normalized constraints", "Makes the search reproducible."),
            ("Source", "Catalyst CaseRegistration through server policy", "Clarifies provenance."),
            ("Reasoning path", "Validate → bounded read → returned rows → synthesis", "Separates data access from language generation."),
            ("Limitations", "Row cap, missing fields, no results, role restriction", "Makes uncertainty and coverage visible."),
        ],
        [1.4, 2.8, 2.8],
    )
    g.h2("28. Case-scoped Sarvam conversation")
    g.p(
        "Talk About This Case is deliberately narrower than database-wide voice Q&A. The user first loads one FIR. The server "
        "retrieves only fields permitted for that role and constructs a system prompt that says the FIR and conversation are "
        "untrusted evidence, not instructions. It forbids general knowledge, other cases, assumptions, invented details, guilt "
        "inference, and identification of unknown people. If the requested fact is absent, the model must use a fixed unavailable "
        "answer. Only eight recent messages are retained, each bounded in length."
    )
    g.table(
        ["Stage", "Model / mechanism", "Input boundary", "Output"],
        [
            ("Listen", "Saaras v3", "Uploaded audio from the current user", "English/Kannada transcript."),
            ("Reason", "Sarvam-30B", "Selected FIR fields + up to 8 recent turns + current question", "Concise grounded response with source field when possible."),
            ("Speak", "Bulbul v3", "Generated case answer", "Audio response."),
        ],
        [1.0, 1.45, 3.1, 1.45],
    )
    g.h2("29. Statement interrogation")
    g.p(
        "The statement-analysis helper tokenizes the question, ranks exact statement sentences by overlapping meaningful terms, "
        "and returns up to three evidence candidates. Groq then answers within that selected statement context. This fixed the "
        "failure mode where a model overlooked negative facts such as “their faces were not visible.” The system should still "
        "display the supporting excerpt so the user can verify it."
    )
    g.h2("30. Similar-case scoring")
    g.table(
        ["Signal", "Points", "Notes"],
        [
            ("Same crime type", "35", "Structured classification match."),
            ("Same pincode", "25", "Exact pincode match."),
            ("Same division", "15", "Normalized structured field match."),
            ("Same status", "5", "Small contextual signal."),
            ("Registered within 90 days", "10", "Or 5 points within one year."),
            ("Same accused name", "40", "Only for investigation-capable roles; identity still requires verification."),
            ("Shared statement terms", "Up to 20", "At least two meaningful shared terms; 4 points each, capped."),
        ],
        [2.0, 0.8, 4.2],
    )
    g.p(
        "Scores are capped at 100; candidates below 20 are omitted and at most eight are returned. This is weighted matching, "
        "not machine learning, and it measures record similarity rather than offender identity, guilt, or legal relevance."
    )
    g.h2("31. Offender review-priority scoring")
    g.table(
        ["Factor", "Maximum / rule", "Interpretation"],
        [
            ("Repeated recorded associations", "Up to 40; 20 per case beyond the first", "More FIRs share the normalized accused name."),
            ("Active-case workload", "15 for 2+, 5 for 1", "Associated cases not marked closed/disposed."),
            ("Recency", "20 within 180 days; 10 within 365", "Recent record association."),
            ("Geographic spread", "Up to 15; 5 per division beyond first", "Cross-division recorded breadth."),
            ("Crime-type breadth", "Up to 10; 5 per type beyond first", "Different recorded classifications."),
        ],
        [2.0, 2.2, 2.8],
    )
    g.p(
        "Priority labels are HIGH at 70+, MEDIUM at 40–69, and STANDARD below 40. The profile includes contributing FIRs, "
        "modus indicators, and identity warnings. No victim demographic factor is used. A same-name group can combine different "
        "people; a shared mobile is only a recorded identifier and still requires ownership verification."
    )
    g.h2("32. Criminal-network analysis")
    g.p(
        "The network begins with normalized accused names. Crucially, an accused-to-accused edge is created only when the two "
        "names are co-recorded in an FIR or share a recorded contact reference. Soft similarity—location, crime type, or "
        "statement-derived modus indicator—may strengthen an existing edge but cannot create one."
    )
    g.table(
        ["Edge component", "Points"],
        [
            ("Co-accused FIRs", "30 each, capped at 60"),
            ("At least one shared recorded mobile", "25"),
            ("Shared location", "5"),
            ("Shared crime classification", "5"),
            ("Shared modus indicators", "5 each, capped at 10"),
        ],
        [3.6, 3.4],
    )
    g.p(
        "Users select a minimum association score of 30, 50, or 70. Connected components become “possible networks for review.” "
        "The interface must never call them proven gangs: coordination and common intent require corroborating evidence."
    )
    g.h2("33. Case brief and investigative leads")
    g.p(
        "The case brief is extractive, not generative. It builds an overview from structured fields, lists source-tagged facts, "
        "extracts up to ten sentences containing explicit dates/times or sequence language, and detects statement excerpts for "
        "CCTV, witnesses, vehicles, digital/telecom, financial references, weapons, and documents/property. Each category maps to "
        "a bounded suggested check such as preserving footage or verifying a transaction through authorized process."
    )
    g.h2("34. Trend and sociological analytics")
    g.p(
        "Trend Analytics counts records inside a selected 6, 12, or 24-month window and reports monthly volume, top crime types, "
        "top divisions, age bands, and data completeness. Social Insights adds crime-by-age composition and a recorded-share "
        "signal only when a crime type has at least three age-known cases, an age band has at least two cases, representation is "
        "at least 1.5× the overall share, and the absolute difference is at least ten percentage points."
    )
    g.callout(
        "Causality guardrail",
        "A composition difference is not a cause, vulnerability verdict, or individual prediction. Gender, occupation, income/economic stress, education, migration, urbanization, and social-category indicators are currently unavailable. Claims about those factors are not implemented.",
        fill=CYAN,
    )
    g.h2("35. Spatial intelligence")
    g.numbered(
        [
            "Use valid latitude/longitude stored on the record when available.",
            "If a record lacks coordinates but another record with the same pincode has coordinates, use that pincode centroid.",
            "If neither exists, use one of five explicit Bengaluru pincode fallbacks.",
            "Otherwise mark the record unmapped; expose mapping source and coverage counts.",
            "Group mapped records by pincode or a rounded coordinate cell and display crime breakdown and contributing FIRs.",
            "For alerts, compare the latest 30 days with the preceding 30 days for one area and crime type.",
        ]
    )
    g.p(
        "An alert requires at least three current-window cases. It is a RECENT_SURGE when the current count is at least twice "
        "the previous count; otherwise it is a REPEAT_CONCENTRATION. Severity is HIGH for an acceleration or at least five "
        "current cases, otherwise MEDIUM. This describes recorded reporting concentration, not certainty of future crime."
    )
    g.h2("36. Crime forecasting")
    g.table(
        ["Component", "Implementation"],
        [
            ("History window", "12 or 24 complete months ending before the current month."),
            ("Recent baseline", "Weighted average of the latest three months with weights 1, 2, and 3."),
            ("Trend", "Linear regression slope over the latest six monthly counts."),
            ("Seasonality", "With 24 months, a bounded 0.5–1.5 month-of-year factor contributes 25% of the adjustment."),
            ("Uncertainty", "At least ±1 case, otherwise 1.64 × residual RMSE from rolling backtests."),
            ("Sufficiency", "INSUFFICIENT below 6 cases or 3 active months; ADEQUATE at 40+ cases and 12+ active months; otherwise LIMITED."),
            ("Output", "Three aggregate monthly estimates with lower/upper range and diagnostics."),
        ],
        [1.55, 5.45],
    )
    g.p(
        "This transparent baseline is appropriate for demonstrating forecast mechanics. It is not a validated crime-prediction "
        "model, is sensitive to reporting changes and sparse data, and must not be used for individual-level predictions."
    )
    g.h2("37. Early Warning Center")
    g.p(
        "The rules engine combines outputs from spatial alerts, offender review profiles, possible networks, and aggregate forecasts. "
        "It sorts HIGH before MEDIUM and retains exact trigger logic, contributing FIRs, recommended checks, and a limitation."
    )
    g.table(
        ["Category", "Minimum trigger"],
        [
            ("SPATIAL_SURGE", "A qualifying spatial concentration/surge alert."),
            ("REPEAT_ASSOCIATION", "3+ associated FIRs, at least one active, latest within 180 days, priority score 60+."),
            ("NETWORK_ACTIVITY", "Possible network score 50+ and a connected FIR within 180 days."),
            ("FORECAST_RISE", "Sufficient data, six-month slope ≥0.25, and next forecast not below recent baseline."),
        ],
        [2.0, 5.0],
    )
    g.h2("38. Guardrails implemented today")
    g.bullets(
        [
            "Provider credentials reside in server environment variables, not browser code.",
            "Authentication and role are verified server-side; the client role is not authoritative.",
            "Passwords use scrypt with random salts.",
            "Sessions are HMAC-signed and expire.",
            "Query fields and filters are role-allowlisted.",
            "Language-model output cannot execute database code.",
            "Text inputs, result counts, audio, history, and requested periods are bounded.",
            "Case conversation is restricted to one selected FIR and treats evidence text as untrusted.",
            "Analytical scores expose factors and non-guilt/non-causal warnings.",
            "Audit records are integrity-signed and omit credentials, full queries, and full victim statements.",
            "Synthetic seed operations are idempotent, labeled, and supervisor-restricted.",
        ]
    )
    g.h2("39. Evaluations that should be added")
    g.table(
        ["Evaluation", "Example metric", "Suggested target before pilot"],
        [
            ("NL query intent accuracy", "Correct fields/filters/sort/limit on a gold set", "≥95% exact policy-valid intent; 100% rejection of malicious cases."),
            ("Answer groundedness", "Every factual clause supported by returned row/field", "≥98%; zero invented FIR numbers."),
            ("Case Q&A faithfulness", "Answer-entailment against selected FIR", "≥98% with explicit abstention on absent facts."),
            ("Kannada quality", "Bilingual expert rating for meaning and naturalness", "Defined acceptance rubric; no safety-significant translation loss."),
            ("Speech recognition", "Word/error rate by dialect, noise, gender, device", "Context-specific baseline and critical-entity accuracy target."),
            ("TTS intelligibility", "Human comprehension of numbers/names/locations", "≥95% critical entity comprehension."),
            ("Network precision", "Sampled edges verified as supported by displayed records", "100% evidence validity; identity verification remains human."),
            ("Forecast backtest", "MAE/RMSE versus naive baselines by division/type", "Must outperform selected baseline before operational use."),
            ("Fairness review", "Error/coverage differences across available groups/regions", "Documented and reviewed; no demographic risk scoring."),
            ("Security red team", "Prompt injection, privilege escalation, data exfiltration", "No unauthorized field or cross-case leakage."),
            ("Performance", "P50/P95 latency and concurrency", "Targets defined from operational workflow."),
        ],
        [2.0, 3.0, 2.0],
        font_size=7.9,
    )


def add_security_testing_deploy(g: Guide) -> None:
    g.h1("Part VII — Security, testing, and deployment", page_break=True)
    g.h2("40. Role and field access")
    g.table(
        ["Role", "Conversational field profile", "Sensitive intelligence", "Audit"],
        [
            ("Constable", "CrimeNo, type, pincode, division, status, registered date", "No repeat sensitive signals/network/profiles", "No"),
            ("Investigator", "Basic + victim name/age/address, accused name/age, statement", "Yes", "No"),
            ("Analyst", "All available fields", "Yes", "No"),
            ("Supervisor", "All available fields", "Yes", "Yes"),
            ("Argos", "All available fields", "Yes; full demo access", "Yes"),
        ],
        [1.2, 2.7, 2.0, 1.1],
    )
    g.h2("41. Important security limitations")
    g.bullets(
        [
            "Registration uses a shared four-digit code (default 3024) and lets the registrant select a role. This is acceptable only for a controlled hackathon; production needs administrator-driven provisioning or enterprise identity with approved role assignment.",
            "Argos has full access and can write cases. A public LinkedIn link therefore exposes persistent data writes and paid AI calls unless sandboxed, rate-limited, and isolated.",
            "Tokens are stored in localStorage, where successful cross-site scripting could steal them. Prefer secure HttpOnly, SameSite cookies or a managed identity flow.",
            "The token format is custom HMAC logic. Production should use a reviewed identity provider, key rotation, revocation, multi-factor authentication, and short-lived access/refresh patterns.",
            "No visible rate limiter, CAPTCHA, abuse quota, IP reputation, or provider-cost circuit breaker exists.",
            "No formal Content Security Policy, security headers, WAF policy, dependency scanning, SAST/DAST, or penetration-test evidence is present.",
            "Third-party AI processing requires contracts, data minimization, retention controls, data-residency review, and explicit authorization.",
            "Audit events share a mutable application table and use a shared HMAC secret. Production should use append-only/WORM storage, key management, retention, export, and independent monitoring.",
            "Sensitive fields are not documented as encrypted at field level; access control alone is not a complete protection strategy.",
            "No formal backup, restoration, disaster recovery, incident response, or breach-notification process is defined.",
        ]
    )
    g.h2("42. Audit trail behavior")
    g.p(
        "Server routes record operational actions with actor, role, action, outcome, target type/ID, bounded safe details, event time, "
        "and an HMAC signature. The reader recomputes the signature to mark tampered content. Reads are limited to 1–200 results "
        "and restricted to Supervisor or Argos. Audit write failure is non-blocking for the protected business action, so the system "
        "can remain available but must expose and monitor `auditUnavailable` conditions in production."
    )
    g.callout(
        "Audit lesson",
        "A signature detects content modification only while the signing key remains secret. It does not make a mutable table append-only, prevent deletion, guarantee clock correctness, or replace external log monitoring.",
        fill=CYAN,
    )
    g.h2("43. Automated test status")
    g.table(
        ["Suite", "Measured result on 26 July 2026", "Coverage themes"],
        [
            ("Backend Node tests", "86 passed, 0 failed", "Auth, safe query policy, FIR validation, conversation history, statements, briefs, explainability, analytics, synthetic data, network, alerts, audit."),
            ("Frontend Vitest", "23 passed across 6 files, 0 failed", "App routing/auth shell, Dashboard, Chat, Deep Dive, Audit Trail, Offender Profiles."),
            ("Total", "109 passed, 0 failed", "A strong unit/component baseline for a hackathon snapshot."),
        ],
        [1.65, 2.3, 3.05],
    )
    g.h2("44. What the current tests do not prove")
    g.bullets(
        [
            "No browser-level end-to-end suite proves login → query → Deep Dive → report → audit in a deployed environment.",
            "No contract tests verify the current Groq, Sarvam, and Catalyst APIs under network failure or schema drift.",
            "No load test establishes concurrency, cold-start latency, provider quota behavior, or cost.",
            "No security test suite covers XSS, CSRF, token theft, brute-force login, role escalation, prompt injection at scale, or denial of service.",
            "No data-quality test is run against a representative real police dataset.",
            "No formal bilingual accuracy, accessibility, usability, or criminology-domain evaluation exists.",
            "Passing tests validate programmed expectations; they do not certify lawful or safe operational use.",
        ]
    )
    g.h2("45. Local development")
    g.code("cd frontend\nnpm install\nnpm run start")
    g.code("cd functions/ks_intelli_pol_function\nnpm install\n# Configure a local .env from .env.example\nnpm test")
    g.p(
        "The frontend API base is `/server/ks_intelli_pol_function`. For an integrated local experience, use Catalyst local serving "
        "or a reverse proxy that preserves this path. Running Vite alone without the function will render the UI but API calls will fail."
    )
    g.h2("46. Required active environment variables")
    g.table(
        ["Variable", "Purpose", "Handling"],
        [
            ("GROQ_API_KEY", "Groq intent and answer/interrogation calls", "Rotate if exposed; store as encrypted Catalyst function variable."),
            ("SARVAM_API_KEY", "Speech and case-conversation APIs", "Server-only encrypted variable; apply quota monitoring."),
            ("AUTH_TOKEN_SECRET", "Session signing, pseudonymous IDs, audit integrity", "At least 32 random characters; separate keys per environment and rotate carefully."),
            ("ATHENA_REGISTRATION_CODE", "Shared demo registration gate", "Defaults to 3024; replace with proper provisioning before any real use."),
        ],
        [2.0, 2.7, 2.3],
    )
    g.h2("47. Build and deployment")
    g.numbered(
        [
            "Run backend tests: `npm test` in `functions/ks_intelli_pol_function`.",
            "Run frontend tests: `npm test` in `frontend`.",
            "Build the frontend: `npm run build` in `frontend`; this runs TypeScript checking and Vite bundling.",
            "Confirm Catalyst project and environment variables; never embed keys in source.",
            "Deploy from the repository root with the Catalyst CLI. `catalyst.json` deploys `frontend/build` and the `ks_intelli_pol_function` target.",
            "Smoke-test registration/login/demo, Dashboard, conversational query, Deep Dive, voice, report, analytics, FIR write, and Audit Trail.",
            "Review Catalyst logs and provider dashboards for errors, latency, and quota usage.",
        ]
    )
    g.h3("Current development URL")
    g.code("https://ks-intellipol-60077550150.development.catalystserverless.in/app/index.html")
    g.p(
        "Treat this as a development/demo endpoint. A production release needs a production Catalyst environment, custom domain and "
        "TLS policy, environment separation, release approvals, rollback, monitoring, and security/governance sign-off."
    )


def add_challenge_status(g: Guide) -> None:
    g.h1("Part VIII — Challenge coverage and production roadmap", page_break=True)
    g.h2("48. Coverage against the 10 expected solution areas")
    rows = [
        ("1. Conversational interface", "LARGELY IMPLEMENTED", "English/Kannada typed Q&A, context, FIR retrieval, voice input, case voice response, local PDF. Bilingual quality needs formal evaluation."),
        ("2. Criminal network analysis", "IMPLEMENTED FOR RECORDED LINKS", "Evidence-qualified accused/case graph and possible groups. No external financial/device/location feeds or entity-resolution service."),
        ("3. Pattern and trend analytics", "IMPLEMENTED AT PROTOTYPE SCALE", "Time, crime type, division, age, hotspots, clusters, seasonal forecast factor. Fixed 300-row cap and limited geocoding."),
        ("4. Sociological insights", "PARTIAL", "Age/crime/geography composition and non-causal share signals. External social/economic indicators and most demographics are absent."),
        ("5. Offender profiling", "IMPLEMENTED AS REVIEW PRIORITY", "Repeat recorded name groups, modus indicators, recency/breadth score. Not a behavioral diagnosis or guilt prediction."),
        ("6. Investigator decision support", "IMPLEMENTED", "Case summaries, timelines, evidence checks, similar cases, and case-scoped Q&A. Outcomes/lead effectiveness are not learned."),
        ("7. Financial crime links", "PARTIAL / NOT INTEGRATED", "Statements may mention accounts/UPI and network methods can surface them; no transaction ingestion, money-trail graph, suspicious transaction engine, or workflow integration."),
        ("8. Forecasting and early warning", "IMPLEMENTED AS TRANSPARENT BASELINE", "Aggregate forecast, spatial surge, repeat association, network, and rising-volume alerts. Not operationally validated prediction."),
        ("9. Explainable AI", "STRONG PROTOTYPE", "Uniform evidence panel, FIR/field citations, rule breakdowns, coverage, limitations. Formal XAI validation and immutable evidence lineage remain."),
        ("10. Secure RBAC/governance", "HACKATHON IMPLEMENTATION", "Signed sessions, field policy, password hashing, audit. Shared role code, full Argos write access, localStorage, and missing enterprise controls prevent production claim."),
    ]
    g.table(["Challenge area", "Status", "Honest assessment"], rows, [2.15, 1.65, 3.2], font_size=7.7)
    g.h2("49. What is production-ready versus demo-ready")
    g.table(
        ["Statement", "Verdict", "Reason"],
        [
            ("“The project is ready for a hackathon demo.”", "YES", "Integrated workflows, synthetic dataset, coherent UI, and passing tests."),
            ("“The project is ready for unrestricted public demo access.”", "ONLY AFTER CONTROLS", "Argos can write and trigger provider costs; isolate data, add quotas/rate limits/CAPTCHA, and monitor abuse."),
            ("“The project is ready for a police pilot with real PII.”", "NO", "Identity, privacy, security, governance, data quality, model evaluation, and operating procedures are incomplete."),
            ("“The system predicts criminals.”", "NO", "Forecasts are aggregate; profiles are review priorities; all outputs require human verification."),
            ("“The system uses fine-tuned Sarvam.”", "NO", "It uses case-scoped prompt grounding with Sarvam-30B."),
        ],
        [3.05, 1.35, 2.6],
    )
    g.h2("50. Phased production roadmap")
    g.h3("Phase 0 — Public-demo hardening")
    g.bullets(
        [
            "Move Argos into an isolated demo tenant/database with automatic record expiry and synthetic-only data.",
            "Add CAPTCHA/bot protection, per-IP and per-session rate limits, provider cost quotas, upload limits, and abuse monitoring.",
            "Remove or sandbox persistent writes; if write demonstrations are required, label and purge them automatically.",
            "Add a visible privacy/demo disclaimer and prohibit real personal information.",
            "Pin dependencies, scan secrets, add security headers/CSP, and fix remaining encoding corruption.",
        ]
    )
    g.h3("Phase 1 — Engineering foundation")
    g.bullets(
        [
            "Define environments, infrastructure/configuration as code, CI/CD checks, approvals, artifact versioning, rollback, and release notes.",
            "Refactor the Advanced I/O monolith into well-owned domain modules or services with structured errors and schemas.",
            "Add runtime validation (for example JSON Schema/Zod), database migrations, indexes, pagination, caching, timeouts, retries, and circuit breakers.",
            "Introduce structured logs, metrics, traces, dashboards, alerts, provider usage/cost telemetry, and synthetic health checks.",
            "Implement backup/restore tests, disaster recovery objectives, and incident response runbooks.",
        ]
    )
    g.h3("Phase 2 — Identity, security, and governance")
    g.bullets(
        [
            "Integrate government/enterprise identity, MFA, device/session management, administrator-assigned roles, least privilege, and joiner/mover/leaver workflows.",
            "Replace localStorage tokens with a reviewed session design and add revocation/key rotation.",
            "Separate append-only audit storage with retention, SIEM export, integrity anchoring, access review, and alerting.",
            "Classify data; define purpose limitation, consent/lawful basis, minimization, retention, deletion, masking, encryption, and break-glass access.",
            "Complete threat modeling, privacy impact assessment, legal review, vendor due diligence, penetration testing, and model risk governance.",
        ]
    )
    g.h3("Phase 3 — Data and analytical reliability")
    g.bullets(
        [
            "Integrate authoritative FIR/case-management interfaces with change data capture, lineage, quality rules, and reconciliation.",
            "Create canonical person, organization, address, device, vehicle, account, and event entities with probabilistic entity resolution and human confirmation.",
            "Use authoritative statewide geocoding and spatial indexes; show geocode accuracy and jurisdiction boundaries.",
            "Build an analytical warehouse/lakehouse or graph store for statewide scale; remove fixed 300-row sampling.",
            "Validate forecast baselines against alternatives; monitor drift and avoid feedback loops from policing intensity.",
        ]
    )
    g.h3("Phase 4 — AI assurance and multilingual quality")
    g.bullets(
        [
            "Build versioned gold datasets for English/Kannada retrieval, follow-ups, case Q&A, abstention, citations, and adversarial prompts.",
            "Add automatic groundedness checks, citation entailment, policy regression, translation evaluation, and human domain review.",
            "Red-team prompt injection embedded inside FIR statements and audio transcripts.",
            "Minimize third-party data, redact when appropriate, and support provider/model fallback with explicit version tracking.",
            "Fine-tune only if evaluations prove prompt grounding or retrieval is inadequate and a lawful, representative training corpus exists.",
        ]
    )
    g.h3("Phase 5 — Advanced features")
    g.bullets(
        [
            "True financial transaction and money-trail ingestion with lawful access controls and temporal graph analytics.",
            "Policy-facing dashboards with aggregate-only privacy controls and scenario analysis.",
            "Event calendars, weather, mobility, census, labor/economic, urbanization, and education indicators with causal caution and provenance.",
            "Collaborative case workspaces, assignments, notes, evidence requests, approval workflows, and integration with existing police systems.",
            "Mobile/offline field experience, alert subscriptions, secure evidence capture, and Kannada-first accessibility.",
            "Validated knowledge graph, cross-jurisdiction federation, and privacy-preserving record linkage.",
        ]
    )


def add_demo_and_learning(g: Guide) -> None:
    g.h1("Part IX — Demonstration, learning, and career lessons", page_break=True)
    g.h2("51. Five-minute demo script")
    g.table(
        ["Time", "Show", "What to say"],
        [
            ("0:00–0:30", "Login / Try Demo", "“Argos creates a short-lived demo session. For production we would isolate and restrict it; today it lets judges explore the full prototype.”"),
            ("0:30–1:00", "Dashboard", "“The dashboard reads live Catalyst records, normalizes jurisdiction labels, and shows selected-division cases rather than hardcoded city cards.”"),
            ("1:00–1:50", "Conversational Hub", "Ask a precise query and a follow-up. “The LLM proposes JSON intent; the server constructs the only executable query.”"),
            ("1:50–2:15", "Evidence panel", "“Every answer exposes FIRs, fields, filters, row bounds, and limitations.”"),
            ("2:15–3:10", "Case Deep Dive", "Load a seeded FIR, show timeline, leads, similar cases, and download the one-page KSP-ATHENA report."),
            ("3:10–3:45", "Talk About This Case", "Ask in Kannada or English and play the spoken answer. “Only the selected role-permitted FIR is in context.”"),
            ("3:45–4:30", "Network / profile / hotspot", "“Edges require recorded co-appearance or contact; similarity alone cannot create an edge.”"),
            ("4:30–5:00", "Early Warnings / close", "“These are explainable human-review signals, not autonomous predictions or findings of guilt.”"),
        ],
        [0.78, 1.42, 4.8],
        font_size=7.8,
    )
    g.h2("52. High-value test questions for the demo")
    g.bullets(
        [
            "“Show vehicle thefts in Koramangala between two dates.”",
            "“Only show the FIR number, case status, and division.”",
            "Follow-up: “Which of those are still under investigation?”",
            "Load a case whose statement says faces were hidden; ask: “Were the suspects’ faces visible?”",
            "Ask a fact that is not in the FIR and verify the case assistant abstains.",
            "Switch the case conversation to Kannada, speak a question, and verify transcript plus audio answer.",
            "Choose a non-Bengaluru division on Dashboard and confirm the cases list appears below the breakdown.",
            "Inspect a network edge and read the exact co-accused/shared-contact evidence.",
            "Select an offender profile and explain each point in its review score.",
            "Filter Crime Forecast and show its sufficiency, backtest MAE, trend, and uncertainty range.",
            "Open Audit Trail as Supervisor/Argos after performing actions and verify recorded events.",
        ]
    )
    g.h2("53. Common demo failure recovery")
    g.table(
        ["Failure", "Calm explanation", "Recovery"],
        [
            ("Voice API delay", "Third-party speech service latency is variable.", "Use typed case Q&A; explain the same scoped pipeline."),
            ("No records found", "The filter may not match the loaded date/data scope.", "Use a known seeded FIR or broaden dates; inspect evidence status."),
            ("Sparse forecast", "The system intentionally refuses to forecast insufficient data.", "Present abstention as a safety feature; choose All Divisions."),
            ("Map has unmapped rows", "Coordinates are incomplete and mapping provenance is explicit.", "Show coverage counts and a directly geocoded seeded region."),
            ("Audit event absent", "Audit writes are non-blocking and storage may be delayed/unavailable.", "Refresh, remove filters, then explain why production monitoring is required."),
            ("AI key/quota error", "Provider configuration/quota is external to the deterministic modules.", "Show Dashboard, Deep Dive brief, network, trends, or forecasts."),
        ],
        [1.5, 3.15, 2.35],
    )
    g.h2("54. Career lessons embedded in the project")
    learning = [
        ("Architecture", "Follow deployed entry points, not folder names. Draw trust boundaries before features."),
        ("API design", "Treat every input as hostile, bound work, and return explicit errors and coverage."),
        ("AI engineering", "LLMs should propose; deterministic policy should authorize. Ground answers and make abstention observable."),
        ("Data engineering", "Names and addresses require normalization, provenance, quality measures, and identity uncertainty."),
        ("Analytics", "A transparent baseline with honest limits is better than an opaque claim of “AI prediction.”"),
        ("Security", "Authentication is not authorization. Authorization is not governance. Governance is not only an audit log."),
        ("Testing", "Test policy and failure behavior, not only happy-path UI. Passing unit tests do not certify the deployed system."),
        ("Product communication", "Say what is active, partial, or roadmap. Judges trust precise limits more than inflated claims."),
        ("Responsible AI", "Never turn correlation, complaint history, or a name match into guilt, identity, or future-offending certainty."),
        ("Production engineering", "Observability, rollback, incident response, cost controls, and data lifecycle are features."),
    ]
    g.table(["Discipline", "Lesson"], learning, [1.55, 5.45])
    g.h2("55. How to explain the project in interviews")
    g.p(
        "Use a structured story: problem → constraints → design decision → implementation → measurement → limitation → next step. "
        "For example: “Natural-language search risked SQL injection and over-broad access. I changed the design so the model returns "
        "JSON intent, then enforced server-side role field/filter allowlists and a 50-row cap. Unit tests cover raw-SQL rejection, "
        "unauthorized fields, wildcard rejection, escaping, and Argos scope. The remaining work is adversarial end-to-end evaluation "
        "and enterprise identity.” This demonstrates engineering judgment rather than feature listing."
    )


FAQ_GROUPS = {
    "56. Product and problem FAQ": [
        ("What is KSP-ATHENA?", "A serverless, explainable crime-intelligence prototype that combines controlled conversational retrieval, selected-case bilingual voice Q&A, case decision support, analytics, relationship visualization, and governance features."),
        ("What user problem does it solve?", "It lets users express investigative questions in natural language instead of manually translating them into database fields and separate analytical tools."),
        ("Who is the primary user?", "Investigators and analysts are the strongest fit. Constables receive restricted lookup/intake; supervisors receive oversight/audit; a future policymaker role should be aggregate-only."),
        ("What is the strongest differentiator?", "The combination of case-scoped bilingual voice interaction and a server-derived evidence panel, backed by deterministic analytics rather than unsupported AI claims."),
        ("Why not use an ordinary chatbot?", "A general chatbot cannot safely enforce database permissions, prove sources, or distinguish absent data from an invented answer. KSP-ATHENA adds a controlled data plane and evidence contract."),
        ("Why combine many modules?", "Investigators move from discovery to case review to relationship/pattern analysis. One workspace reduces context switching and lets citations connect those stages."),
        ("Is it replacing investigators?", "No. It retrieves, organizes, summarizes, and prioritizes evidence for authorized human review. Legal and operational decisions remain human and follow procedure."),
        ("What does ATHENA stand for?", "The repository uses KSP-ATHENA as the product identity rather than a formally encoded acronym. Avoid inventing an expansion unless the team officially approves one."),
        ("Why is this suitable for a hackathon?", "It demonstrates an end-to-end idea with real integration, meaningful safety design, visible analytics, bilingual voice, synthetic test signals, and measurable automated tests."),
        ("What is the biggest current limitation?", "Production governance: public Argos writes, shared-code role registration, third-party sensitive-data processing, limited data scale/geocoding, and absent operational validation."),
    ],
    "57. Architecture and technology FAQ": [
        ("Why React?", "It supports a modular, interactive single-page workspace with reusable components and a large ecosystem for charts, maps, graphs, icons, and testing."),
        ("Why TypeScript?", "It makes frontend payload contracts and component props easier to reason about and catches many mistakes before bundling."),
        ("Why Vite instead of Create React App?", "Vite provides fast development startup and a modern, simple production build. Create React App is no longer the preferred modern default."),
        ("Why Zoho Catalyst?", "The team already had the project and tables provisioned; Catalyst provides client hosting, a Node function, Data Store, environment variables, and a single deployment workflow."),
        ("Why one Advanced I/O function?", "It minimized hackathon deployment complexity and centralized security policy. The trade-off is a growing monolith that should be modularized for production."),
        ("Why Express?", "It gives familiar routing and middleware inside the Advanced I/O function, making many APIs straightforward to organize."),
        ("Why JavaScript/Node for analytics?", "The current algorithms are bounded deterministic transformations, so keeping them in the function simplifies deployment and testing. Heavy graph/ML workloads may later move to specialized services."),
        ("Why are Python folders present?", "They are earlier prototypes for ML/graph/analytics services. They are not imported or deployed by the active Catalyst configuration."),
        ("Does the project use FastAPI in production?", "No. FastAPI-related files exist only in disconnected prototype directories."),
        ("Does the project use NetworkX or scikit-learn in production?", "No. Active network and forecast logic is implemented in Node modules."),
        ("Why Leaflet?", "It is an established open mapping library and integrates cleanly through React Leaflet."),
        ("Why Recharts?", "It makes declarative charts easy to build from React state and API responses."),
        ("Why a force graph?", "Networks are relational; an interactive node-link view makes accused-to-case and association structure easier to explore than a flat table."),
        ("Why client-side PDF generation?", "It creates immediate local reports without storing another sensitive artifact on the server. Production may need text-native signed reports and controlled records management."),
        ("What is the deployment unit?", "The compiled `frontend/build` client and the `ks_intelli_pol_function` Catalyst function, as declared in `catalyst.json`."),
    ],
    "58. AI and language FAQ": [
        ("Where exactly is generative AI used?", "Groq handles JSON intent generation, grounded conversational synthesis, and statement interrogation; Sarvam handles case-scoped bilingual conversation plus speech-to-text/text-to-speech."),
        ("Does Groq write SQL?", "No. It proposes a small JSON intent. The server validates that object and constructs the only executable ZCQL SELECT."),
        ("Why is direct AI-generated SQL dangerous?", "It can introduce injection, unauthorized fields, destructive commands, excessive scans, or unpredictable syntax. Server construction turns language flexibility into a controlled query."),
        ("How are hallucinations reduced?", "Limit context to returned rows or one FIR, require absence responses, independently build citations, display evidence/coverage, use low temperatures, and keep decisions human."),
        ("Can hallucination be eliminated?", "No. Probabilistic model behavior remains. High assurance requires automated groundedness checks, adversarial evaluations, user verification, and safe fallback/abstention."),
        ("Why use Groq?", "It offers low-latency hosted inference suitable for live conversational demonstrations. Production selection must also consider privacy, availability, cost, contracts, and model lifecycle."),
        ("Why use Sarvam?", "Its India-focused speech and language capabilities suit English/Kannada interaction and complete a two-way voice workflow."),
        ("Is Sarvam fine-tuned on this database?", "No. The current system passes one selected FIR as bounded prompt context. That is grounding, not weight fine-tuning."),
        ("Why scope voice conversation to one case?", "It reduces context size, latency, cost, cross-case leakage, and confusion while matching the investigator’s immediate task."),
        ("How is follow-up context handled?", "The Conversational Hub stores bounded turns by pseudonymous user and session; the case conversation sends only the latest eight sanitized turns."),
        ("How does the assistant answer Kannada?", "The request language determines the prompt language; Sarvam-30B returns Kannada and Bulbul v3 can speak it. Quality still needs native-speaker evaluation."),
        ("What happens if a fact is absent?", "The case assistant is instructed to return a fixed “not available in the selected FIR” response rather than infer."),
        ("Can the FIR statement prompt-inject the model?", "It is a risk. The system prompt explicitly labels FIR and history as untrusted evidence and forbids following their instructions. Red-team evaluations and defense in depth are still required."),
        ("Why not use embeddings/vector search?", "The dataset and use cases were addressed with structured filters, exact evidence, and deterministic similarity for the hackathon. Vector search could help semantic retrieval later, but adds indexing, privacy, evaluation, and citation challenges."),
        ("Would RAG be a correct description?", "The system has retrieval-grounded generation: records are retrieved before answer synthesis. It is not currently a vector-database RAG pipeline."),
        ("What model temperature is used?", "Intent generation uses 0 for consistency; grounded answer synthesis uses 0.3. Deterministic analytics use no model temperature."),
        ("How do you evaluate bilingual AI?", "Use a versioned English/Kannada gold set, native domain reviewers, semantic fidelity, critical entity accuracy, abstention correctness, citation entailment, and speech error/intelligibility metrics."),
        ("Could the model expose another case?", "The case route supplies only the selected role-permitted FIR. The broader chat route can retrieve multiple permitted rows, so role policy, limits, session isolation, and leak tests are essential."),
    ],
    "59. Data and analytics FAQ": [
        ("Are the analytics hardcoded?", "The rules and thresholds are coded, but results are dynamically calculated from database records. Seeded data is designed to exercise those rules."),
        ("Is rule-based the same as hardcoded output?", "No. Hardcoded output would return fixed conclusions. A rule-based system applies explicit general logic to changing input data."),
        ("Is the offender score a prediction of crime?", "No. It is an investigative review-priority score based on recorded associations, active cases, recency, geography, and crime-type breadth."),
        ("Why is name matching risky?", "Different people can share names, spellings vary, and one person may use aliases. The UI warns that identity needs independent identifiers."),
        ("What creates a criminal-network edge?", "At least one direct recorded link: co-appearance in an FIR or a shared recorded contact. Location/crime/modus similarity only strengthens an existing edge."),
        ("Does a network prove an organized gang?", "No. It visualizes recorded associations. Organized coordination and common intent require corroboration."),
        ("How are similar cases scored?", "Crime type 35, pincode 25, division 15, status 5, time proximity 5/10, same accused 40, and shared statement terms up to 20; capped at 100."),
        ("What does the forecast predict?", "Aggregate monthly case volume for the selected division/crime type for three months, with a range and diagnostics."),
        ("Why not forecast when data is sparse?", "A numerical output can look authoritative even when unsupported. The sufficiency rule intentionally abstains below six cases or three active months."),
        ("Is the forecast machine learning?", "It is a transparent statistical baseline: recent weighted average, six-month linear slope, optional seasonal adjustment, and residual uncertainty."),
        ("What is MAE?", "Mean absolute error: the average absolute difference between earlier rolling predictions and actual monthly counts. Lower is better within comparable data."),
        ("What is RMSE used for?", "Root mean square residual error drives the displayed uncertainty range; it penalizes larger errors more strongly."),
        ("What is a hotspot?", "A mapped pincode or coordinate cell with a concentration of recorded cases. It may reflect reporting and data coverage as well as underlying events."),
        ("How are statewide locations mapped?", "Direct coordinates first, then same-pincode centroid, then only five explicit Bengaluru fallbacks; all other missing coordinates remain unmapped."),
        ("What is a spatial surge?", "At least three same-type cases in one area in the latest 30 days and at least twice the preceding 30-day count."),
        ("Does social insight identify causes?", "No. It describes recorded composition. Causal claims require external variables, study design, confounding controls, ethics review, and domain expertise."),
        ("Why a 300-row cap?", "It bounds demo latency and function/database work. It also means analytics may be incomplete and must be replaced with scalable queries or an analytical store."),
        ("Are synthetic records realistic evidence?", "They are scenario-rich test fixtures, not evidence about real crime, model accuracy, bias, or operational benefit."),
        ("How should financial analysis be described?", "Partial. The system can surface financial references and shared indicators, but it lacks transaction feeds, money-trail graphs, suspicious transaction detection, and financial workflow integration."),
        ("How do you prevent feedback loops?", "The current prototype only warns that outputs are not causal. Production must measure how patrol/reporting changes alter data, prevent enforcement intensity from becoming self-confirming evidence, and audit outcomes."),
    ],
    "60. Security and governance FAQ": [
        ("How are passwords stored?", "Using Node scrypt with a random per-user salt; plaintext passwords are never stored."),
        ("What is AUTH_TOKEN_SECRET?", "A high-entropy server secret used to sign sessions, derive pseudonymous IDs, and sign audit events. It must be at least 32 random characters and never committed."),
        ("Why is one secret for several purposes not ideal?", "Key separation limits blast radius and simplifies rotation. Production should use separate managed keys for sessions, pseudonyms, and audit integrity."),
        ("What does role-based access control enforce?", "Server-side permitted fields, filters, and selected sensitive modules. The client menu is only a usability layer."),
        ("Can users choose their role?", "Currently yes during registration if they know the shared code. That is a hackathon shortcut and unacceptable for production."),
        ("Why is the registration code 3024?", "It is the requested default demo code, configurable by environment variable. It is not a strong authentication factor."),
        ("Is Argos safe for LinkedIn public access?", "Not as currently privileged without isolation. It can write cases and invoke paid APIs. Add synthetic-only tenant isolation, expiry, rate limits, CAPTCHA, quotas, and monitoring."),
        ("Why is localStorage a concern?", "JavaScript can read it; an XSS vulnerability could steal the token. Secure HttpOnly cookies or managed identity reduce this exposure."),
        ("What is recorded in the audit trail?", "Actor pseudonym/role, action, outcome, target, bounded safe details, timestamp, and HMAC signature—not passwords, tokens, full chat queries, or full statements."),
        ("Can audit signatures prevent deletion?", "No. They detect modified content if the key is safe; append-only storage, retention controls, independent export, and monitoring are still needed."),
        ("Why are audit failures non-blocking?", "Availability of the main action was prioritized for the prototype. In production, critical actions may require fail-closed behavior or prominent monitoring depending on policy."),
        ("Is the system compliant with law-enforcement regulation?", "No compliance certification is claimed. Legal basis, privacy impact, security standards, evidence rules, retention, vendor contracts, and accountability require formal review."),
        ("How should secrets be deployed?", "As encrypted Catalyst function environment variables, separately for development/staging/production, with least access, rotation, and leak scanning."),
        ("What data goes to external AI providers?", "Groq receives bounded query context/returned fields for conversational tasks; Sarvam receives selected case context and/or audio. Production needs strict minimization, contracts, and residency/retention review."),
        ("What is the most important public-demo control?", "Prevent real personal data entry and isolate/purge all demo writes while enforcing rate and cost limits."),
        ("How would you handle a data breach?", "The project does not yet define this. Production needs detection, containment, key/token revocation, evidence preservation, impact assessment, notification procedure, recovery, and post-incident review."),
        ("What is least privilege?", "Give each user and service only the data and actions required for the current purpose, for the minimum duration."),
        ("What is a break-glass role?", "A tightly controlled emergency override with justification, short duration, approval/alerting, enhanced audit, and retrospective review."),
    ],
    "61. Testing, deployment, and operations FAQ": [
        ("How many automated tests pass?", "109 total in the documented snapshot: 86 backend and 23 frontend tests."),
        ("What is the strongest test area?", "Backend policy/rule logic: safe queries, auth, explainability, analytics thresholds, seed integrity, and audit behavior."),
        ("What is the biggest testing gap?", "Deployed end-to-end, adversarial security, provider contract, performance, bilingual quality, accessibility, and real-data validation."),
        ("What happens when Groq is unavailable?", "AI-dependent chat/interrogation fails, but deterministic dashboards and analytics can still function. Production should add explicit timeout/retry/circuit-breaker behavior and user-friendly fallback."),
        ("What happens when Sarvam is unavailable?", "Voice and case conversational response fail; typed deterministic case brief/deep-dive data remains available."),
        ("What is a cold start?", "A serverless function may need to initialize after inactivity, increasing the first request’s latency."),
        ("How should latency be measured?", "Capture end-to-end and per-stage P50/P95/P99: authentication, database, intent model, policy, answer model, STT, TTS, and client rendering."),
        ("How should cost be controlled?", "Per-user/session quotas, rate limits, bounded context, caching where safe, provider budgets/alerts, short demo expiry, and model routing."),
        ("Why pin dependencies?", "Using fixed/locked versions improves reproducibility and prevents an unexpected upstream change from entering a deployment."),
        ("What is CI/CD?", "Automated build, test, security, packaging, approval, deployment, smoke test, and rollback workflow for every change."),
        ("How would you roll back?", "Deploy versioned artifacts/configuration and retain a known-good release. Current documentation does not yet define an automated rollback process."),
        ("What should be monitored?", "Availability, latency, errors, 401/403/429 rates, audit failures, database capacity, model/provider failures, cost/quota, suspicious access, and data-quality drift."),
        ("What is an SLO?", "A measurable reliability objective, such as 99.9% availability or a P95 chat latency target, supported by error budgets."),
        ("Why separate dev and production?", "To prevent test data/keys/experiments from affecting real users and to enable controlled promotion and rollback."),
        ("How should schema changes be handled?", "Versioned migrations, compatibility tests, backups, staged rollout, and explicit ownership—not manual undocumented edits."),
    ],
    "62. Ethics, criminology, and responsible-use FAQ": [
        ("Why must the system avoid guilt language?", "FIRs contain allegations and preliminary information. Legal guilt is determined through due process, not a database association or model score."),
        ("Can demographic data be used for offender risk?", "The current design does not do so. Any demographic use needs legal basis, necessity/proportionality, fairness analysis, and strict prohibition of discriminatory profiling."),
        ("What is correlation versus causation?", "Correlation means variables vary together; causation means one produces change in another. Crime data has many confounders and reporting biases, so descriptive patterns cannot establish cause."),
        ("What is automation bias?", "People may over-trust computer outputs. Evidence panels, uncertainty, training, workflow approvals, and audits reduce but do not eliminate it."),
        ("What is selective-label bias?", "Recorded outcomes reflect who was observed, reported, investigated, or charged—not all real events. Training or evaluation on these labels can reproduce institutional patterns."),
        ("What is a feedback loop?", "More policing can produce more recorded incidents in an area, which a system may then use to justify still more policing."),
        ("How does explainability help accountability?", "It lets a reviewer inspect sources, fields, rules, thresholds, and limits. Explainability is necessary but not sufficient; decisions and outcomes must also be governed."),
        ("Should an alert automatically dispatch police?", "No. It should prompt contextual review, data verification, proportional planning, and authorized decision-making."),
        ("How should false positives be handled?", "Track them, show uncertainty, require evidence review, make decisions reversible, provide challenge/correction paths, and tune thresholds through validated evaluation."),
        ("How should false negatives be handled?", "Do not imply complete coverage; expose unmapped/missing/capped data, monitor recall on validated sets, and retain ordinary investigative processes."),
        ("Who should approve production use?", "Police leadership, security, legal/privacy, data governance, domain experts, responsible-AI reviewers, operations, and affected oversight authorities—not the development team alone."),
        ("What documentation should accompany each model?", "Purpose, owner, version, provider, data sent, limitations, evaluations, thresholds, monitoring, incidents, change log, and retirement plan."),
    ],
}


def add_faq(g: Guide) -> None:
    g.h1("Part X — Detailed judge and interview FAQ", page_break=True)
    g.p(
        "Use these answers as a study guide, not a script to memorize word-for-word. A strong response starts with the direct answer, "
        "then gives one implementation fact, one limitation, and one next step."
    )
    for heading, questions in FAQ_GROUPS.items():
        g.h2(heading)
        for question, answer in questions:
            g.qa(question, answer)


def add_appendices(g: Guide) -> None:
    g.h1("Part XI — Appendices", page_break=True)
    g.h2("Appendix A — Repository map")
    g.table(
        ["Path", "Responsibility", "Deployment status"],
        [
            ("frontend/src/App.tsx", "Authentication shell, lazy routing, grouped sidebar, session actions", "ACTIVE"),
            ("frontend/src/components/*", "Feature screens and visualizations", "ACTIVE"),
            ("frontend/src/api.ts", "Authenticated fetch wrapper and API base", "ACTIVE"),
            ("frontend/src/styles/*", "Typography and neo-brutalist design system", "ACTIVE"),
            ("functions/ks_intelli_pol_function/index.js", "Express routes and provider/database orchestration", "ACTIVE"),
            ("functions/ks_intelli_pol_function/query-policy.js", "Role fields/filters and safe ZCQL construction", "ACTIVE"),
            ("functions/ks_intelli_pol_function/auth.js", "Registration, login, demo, session verification", "ACTIVE"),
            ("functions/ks_intelli_pol_function/*-analytics / intelligence modules", "Deterministic feature logic", "ACTIVE"),
            ("database/*.schema.md", "Manually provisioned Catalyst table documentation", "REFERENCE"),
            ("database/seed/*", "Synthetic datasets and export/check utilities", "ACTIVE DATA / ADMIN"),
            ("config/karnataka-divisions.json", "Pincode-to-division reference data", "REFERENCE; not yet used by FIR insert"),
            ("backend/*", "Earlier agent/orchestration/Python service prototypes", "DORMANT"),
            ("catalyst.json", "Client and function deployment targets", "ACTIVE"),
        ],
        [2.65, 3.05, 1.3],
        font_size=7.7,
    )
    g.h2("Appendix B — Key configuration facts")
    g.table(
        ["Item", "Value / behavior"],
        [
            ("Catalyst project", "KS-IntelliPol (development environment in documented snapshot)"),
            ("Function target", "ks_intelli_pol_function"),
            ("Client source", "frontend/build"),
            ("API base", "/server/ks_intelli_pol_function"),
            ("Normal session", "8 hours"),
            ("Argos session", "30 minutes"),
            ("Chat row limit", "1–50; default intent limit 20"),
            ("Conversation context", "12 recent stored messages by default; loader maximum 30"),
            ("Case conversation history", "8 recent messages; 1,000 characters each"),
            ("Case question maximum", "500 characters"),
            ("Analytics retrieval cap", "300 records for most cross-record modules"),
            ("Audit read limit", "1–200; underlying read bounded at 200"),
            ("Trend windows", "6, 12, or 24 months"),
            ("Forecast history", "12 or 24 months; predicts 3 months"),
            ("Network thresholds", "30, 50, or 70"),
            ("Supported FIR crime types", "Theft, Assault, Cyber Crime, Fraud, Missing Person"),
            ("Synthetic sets", "25 V1 + 30 V2, clearly labeled"),
        ],
        [2.45, 4.55],
    )
    g.h2("Appendix C — Glossary")
    glossary = [
        ("API", "A defined way for software components to request data or actions from one another."),
        ("Authentication", "Proving who a user is."),
        ("Authorization", "Deciding what an authenticated user may do or see."),
        ("RBAC", "Role-based access control: permissions assigned through roles."),
        ("LLM", "Large language model; a probabilistic model that generates/understands text."),
        ("Grounding", "Constraining an answer to supplied evidence such as retrieved FIR fields."),
        ("Fine-tuning", "Updating model weights using a training dataset; not currently used for Sarvam in this project."),
        ("RAG", "Retrieval-augmented generation: retrieve evidence before generating an answer."),
        ("Prompt injection", "Malicious or accidental instructions inside untrusted content that try to override system rules."),
        ("Hallucination", "A fluent but unsupported or incorrect model output."),
        ("ZCQL", "Zoho Catalyst Query Language used to read Data Store tables."),
        ("Serverless", "Cloud execution where infrastructure scaling/runtime management is largely delegated to the platform."),
        ("HMAC", "Keyed cryptographic message authentication code used to verify integrity/authenticity."),
        ("scrypt", "Memory-hard password derivation function designed to resist brute-force attacks."),
        ("Pseudonymization", "Replacing a direct identifier with a stable derived value; it is not the same as anonymization."),
        ("Entity resolution", "Deciding whether records refer to the same real-world person/object, with uncertainty."),
        ("Hotspot", "An area with a concentration of recorded events under a stated method/window."),
        ("MAE", "Mean absolute error, an average magnitude of forecast errors."),
        ("RMSE", "Root mean square error, which gives larger errors more influence."),
        ("Confidence/uncertainty range", "A displayed range around a point estimate; interpretation depends on the method and assumptions."),
        ("Data lineage", "Trace of where data came from, how it changed, and which output used it."),
        ("Observability", "Logs, metrics, and traces that explain a live system’s behavior."),
        ("SLO", "Service level objective: a measurable reliability/performance target."),
        ("WORM", "Write once, read many storage, useful for tamper-resistant audit retention."),
        ("Human in the loop", "A person reviews and owns decisions rather than allowing full automation."),
    ]
    g.table(["Term", "Plain-language meaning"], glossary, [1.7, 5.3], font_size=8.15)
    g.h2("Appendix D — Pre-demo checklist")
    g.bullets(
        [
            "Confirm development URL loads in a clean/private browser.",
            "Confirm Try Demo creates Argos and the 30-minute expiry is acceptable.",
            "Verify no real personal information is present.",
            "Run 109 automated tests and a production frontend build.",
            "Check Groq and Sarvam keys, quotas, and Catalyst logs.",
            "Preselect one reliable query and two FIR numbers for Deep Dive.",
            "Test microphone permission, Kannada transcript, TTS volume, and fallback typed question.",
            "Verify one selected division lists basic cases.",
            "Generate one-page KSP-ATHENA report and conversation PDF.",
            "Perform an auditable action, refresh Audit Trail, and remove restrictive filters.",
            "Prepare an offline screenshot/video fallback in case venue internet fails.",
            "State the boundaries: synthetic data, human review, no guilt prediction, not yet production-ready.",
        ]
    )
    g.h2("Appendix E — Production gate checklist")
    gates = [
        ("Identity", "Enterprise identity, MFA, admin role provisioning, revocation, access review"),
        ("Security", "Threat model, pen test, CSP/WAF/rate limits, secrets/KMS, dependency and code scans"),
        ("Privacy", "Legal basis, DPIA, minimization, retention, residency, vendor contracts, subject/correction process"),
        ("Data", "Authoritative integration, lineage, quality SLAs, entity resolution, geocoding, schema/version governance"),
        ("AI", "Gold sets, groundedness, bilingual/speech quality, adversarial testing, drift/version monitoring"),
        ("Analytics", "Backtesting, baseline comparison, bias/feedback-loop study, threshold governance"),
        ("Operations", "CI/CD, staged environments, rollback, observability, on-call, incident response, backup/DR"),
        ("Audit", "Append-only store, key separation, SIEM export, retention, deletion detection"),
        ("UX/accessibility", "WCAG testing, keyboard/screen reader, Kannada typography, mobile/responsive, training"),
        ("Governance", "Owners, approvals, model/data cards, standard operating procedures, change and incident review"),
    ]
    g.table(["Gate", "Required evidence before real operational use"], gates, [1.45, 5.55])
    g.h2("Closing perspective")
    g.p(
        "KSP-ATHENA’s strongest achievement is not that it places “AI” beside police data. It is that the prototype begins to "
        "separate flexible language from controlled authority, exposes evidence behind outputs, and frames analytics as human-review "
        "signals. The next stage is less about adding buttons and more about proving reliability, protecting people, integrating "
        "authoritative systems, and building accountable operations. That transition—from a compelling demo to a trustworthy "
        "socio-technical system—is the central professional lesson of the project."
    )


def build_document() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    architecture_path = ASSET_DIR / "architecture.png"
    query_path = ASSET_DIR / "query_pipeline.png"
    analytics_path = ASSET_DIR / "analytics_pipeline.png"
    make_architecture_diagram(architecture_path)
    make_query_pipeline(query_path)
    make_analytics_diagram(analytics_path)

    doc = Document()
    configure_styles(doc)
    doc.core_properties.title = "KSP-ATHENA Complete Project Guide"
    doc.core_properties.subject = "Architecture, technology stack, AI safety, analytics, security, deployment, FAQ, and production roadmap"
    doc.core_properties.author = "KSP-ATHENA Project Team"
    doc.core_properties.keywords = "KSP-ATHENA, crime intelligence, Zoho Catalyst, React, explainable AI, Groq, Sarvam AI"
    doc.core_properties.comments = "Implementation snapshot prepared from the KSP-Athena repository on 26 July 2026."

    add_cover(doc)
    # Main content gets an independent section so title-page header settings do not leak.
    main_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(main_section)
    g = Guide(doc)
    add_front_matter(g)
    add_executive_summary(g)
    add_personas_and_tour(g)
    add_architecture(g, architecture_path, query_path, analytics_path)
    add_stack(g)
    add_data_model(g)
    add_ai_and_analytics(g)
    add_security_testing_deploy(g)
    add_challenge_status(g)
    add_demo_and_learning(g)
    add_faq(g)
    add_appendices(g)
    add_headers_and_footers(doc)

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    output = build_document()
    print(output)
