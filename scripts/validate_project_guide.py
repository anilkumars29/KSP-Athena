from __future__ import annotations

import re
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "docs" / "KSP-ATHENA_Complete_Project_Handbook.docx"


def main() -> None:
    document = Document(DOCX)
    paragraphs = [p for p in document.paragraphs if p.text.strip()]
    headings = [p for p in paragraphs if p.style.name.startswith("Heading")]
    required_headings = [
        "Part I — Executive understanding",
        "Part IV — Technology stack and the reason for every choice",
        "Part VI — AI, analytics, explainability, and guardrails",
        "Part X — Detailed judge and interview FAQ",
        "Part XI — Appendices",
    ]
    missing_headings = [item for item in required_headings if item not in {p.text for p in headings}]

    with zipfile.ZipFile(DOCX) as archive:
        media = [name for name in archive.namelist() if name.startswith("word/media/")]
        document_xml = archive.read("word/document.xml").decode("utf-8")
        core_xml = archive.read("docProps/core.xml").decode("utf-8")
    all_text_parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text_parts.extend(p.text for p in cell.paragraphs)
    all_text = " ".join(all_text_parts)
    words = re.findall(r"\b[\w’'-]+\b", all_text, flags=re.UNICODE)

    alt_descriptions = re.findall(r'<wp:docPr[^>]+descr="([^"]+)"', document_xml)
    page_breaks = len(re.findall(r'w:type="page"', document_xml))
    consecutive_page_breaks = bool(re.search(r'w:type="page".{0,700}w:type="page"', document_xml, flags=re.DOTALL))
    qa_count = sum(1 for p in paragraphs if p.text.startswith("Q. "))
    tables = len(document.tables)

    checks = {
        "file_exists": DOCX.exists(),
        "file_size_over_200kb": DOCX.stat().st_size > 200_000,
        "paragraphs_over_450": len(paragraphs) > 450,
        "word_count_over_10000": len(words) > 10_000,
        "headings_over_50": len(headings) > 50,
        "tables_over_20": tables > 20,
        "faq_over_100": qa_count > 100,
        "three_diagrams": len(media) == 3,
        "all_diagrams_have_alt_text": len(alt_descriptions) == len(media),
        "no_required_heading_missing": not missing_headings,
        "title_metadata_present": "KSP-ATHENA Complete Project Guide" in core_xml,
        "two_sections": len(document.sections) == 2,
        "page_breaks_present": page_breaks >= 10,
        "no_close_consecutive_page_breaks": not consecutive_page_breaks,
    }

    for name, passed in checks.items():
        print(f"{'PASS' if passed else 'FAIL'} {name}")
    print(f"METRIC file_bytes={DOCX.stat().st_size}")
    print(f"METRIC paragraphs={len(paragraphs)}")
    print(f"METRIC words={len(words)}")
    print(f"METRIC headings={len(headings)}")
    print(f"METRIC tables={tables}")
    print(f"METRIC faq_questions={qa_count}")
    print(f"METRIC diagrams={len(media)}")
    print(f"METRIC alt_descriptions={len(alt_descriptions)}")
    print(f"METRIC explicit_page_breaks={page_breaks}")
    if missing_headings:
        print("MISSING headings=" + " | ".join(missing_headings))

    failed = [name for name, passed in checks.items() if not passed]
    if failed:
        raise SystemExit("Validation failed: " + ", ".join(failed))


if __name__ == "__main__":
    main()
